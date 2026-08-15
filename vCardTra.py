#!/usr/bin/env python3
"""
vCardTra — Convert vCard (.vcf) files to text, JSON, and other formats.
Supports vCard 2.1, 3.0, and 4.0.
"""

__version__ = '4.0'

import sys
import re
import glob
import json
import base64
import csv
import io
import html
import sqlite3
from xml.sax.saxutils import escape as xml_escape
from datetime import datetime
from pathlib import Path
import quopri


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

# All known vCard property names across versions 2.1, 3.0, 4.0.
# Anything outside this set is treated as a custom/vendor field.
STANDARD_FIELDS = {
    'FN', 'N', 'TEL', 'EMAIL', 'ADR', 'ADDR', 'ORG', 'TITLE', 'NOTE', 'BDAY',
    'LABEL', 'URL', 'NICKNAME', 'ANNIVERSARY', 'GENDER', 'CATEGORIES',
    'CREATED', 'REVISED', 'VERSION', 'BEGIN', 'END', 'REV', 'PHOTO', 'SOUND',
    'KEY', 'LOGO', 'AGENT', 'PRODID', 'PROFILE', 'SOURCE', 'NAME', 'CLASS',
    'SORT-STRING', 'UID', 'MAILER', 'TZ', 'GEO', 'ROLE', 'MEMBER', 'RELATED',
    'LANG', 'IMPP', 'XML', 'CLIENTPIDMAP', 'CALADRURI', 'CALURI', 'FBURL',
    'EXPERTISE', 'HOBBY', 'INTEREST', 'ORG-DIRECTORY',
}

# Binary or internal fields — silently ignored, never shown as custom.
# PHOTO/LOGO are intentionally NOT here as of v4 — they get real handling
# (decoded to bytes) in parse_vcard(). SOUND/KEY have no practical use in
# any current or planned output format, so they stay silent indefinitely.
SILENT_FIELDS = {
    'VERSION', 'BEGIN', 'END', 'SOUND', 'KEY',
    'AGENT', 'PRODID', 'PROFILE', 'SOURCE', 'NAME', 'CLASS',
}

# Image types recognised on PHOTO/LOGO TYPE parameters (v2.1 inline style,
# e.g. PHOTO;JPEG;ENCODING=BASE64:...). Used by extract_media_type().
KNOWN_IMAGE_TYPES = {'JPEG', 'JPG', 'PNG', 'GIF', 'BMP', 'TIFF', 'WEBP'}

MONTHS = [
    'January', 'February', 'March', 'April', 'May', 'June',
    'July', 'August', 'September', 'October', 'November', 'December',
]

DIVIDER     = '─' * 60
DIVIDER_HVY = '=' * 60
DIVIDER_CTX = '-' * 60

# Fields that produce visible output. Used to decide if a contact
# has any useful content worth writing to the output file.
DISPLAYABLE_FIELDS = {
    'name', 'phones', 'emails', 'addresses', 'organizations',
    'title', 'urls', 'notes', 'birthday', 'nickname', 'label',
    'anniversary', 'gender', 'categories', 'created', 'revised',
    'role', 'timezone', 'geo', 'impp',
}

# Output formats implemented so far, mapped to their default file extension.
# --format validates against this plus PLANNED_FORMATS (see main()).
FORMAT_EXTENSIONS = {
    'text': 'txt',
    'json': 'json',
    'csv': 'csv',
    'markdown': 'md',
    'md': 'md',
    'html': 'html',
    'ics': 'ics',
    'sql': 'sql',
    'sqlite': 'db',
    'xml': 'xml',
    'pdf': 'pdf',
    'docx': 'docx',
}

# Formats speced in CONTEXT.md but not yet implemented — kept separate from
# an "unknown format" typo so the error message can point at what's coming.
PLANNED_FORMATS = set()

# Formats that need an optional third-party library. Checked lazily inside
# each formatter (try/except ImportError) — never required, never
# auto-installed. Listed here so main() can give a specific install hint
# up front rather than waiting for the ImportError deep in the write step.
FORMAT_DEPENDENCIES = {
    'pdf': 'weasyprint',
    'docx': 'python-docx',
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def strip_uri_prefix(value):
    """Remove tel: or mailto: URI prefixes (case-insensitive)."""
    low = value.lower()
    if low.startswith('tel:'):
        return value[4:]
    if low.startswith('mailto:'):
        return value[7:]
    return value


def extract_type(key_part):
    """
    Return the TYPE label for a field (e.g. 'Work/Voice', 'Home').
    Handles v3.0+ TYPE=WORK,VOICE and v2.1 inline TEL;WORK;VOICE.
    Always Title Case.
    """
    match = re.search(r'TYPE=([^;:]+)', key_part, re.IGNORECASE)
    if match:
        parts = [p.strip().title() for p in match.group(1).split(',') if p.strip()]
        return '/'.join(parts) if parts else None

    parts = key_part.split(';')
    if len(parts) > 1:
        ignore = {'INTERNET', 'QUOTED-PRINTABLE', 'BASE64', 'ENCODING', 'CHARSET'}
        params = [
            p.strip().title()
            for p in parts[1:]
            if p.strip() and p.strip().upper() not in ignore
        ]
        return '/'.join(params) if params else None

    return None


def extract_media_type(key_part):
    """
    Return the image type for a PHOTO or LOGO field (e.g. 'JPEG', 'PNG'),
    or None if not specified.
    Handles v3.0+ TYPE=JPEG and v2.1 inline PHOTO;JPEG;ENCODING=BASE64 style
    (checked against KNOWN_IMAGE_TYPES rather than extract_type()'s generic
    ignore-list, since v2.1 params like 'ENCODING=BASE64' aren't cleanly
    filtered by that function's exact-match check).
    """
    match = re.search(r'TYPE=([A-Za-z0-9-]+)', key_part, re.IGNORECASE)
    if match:
        return match.group(1).upper()
    for part in key_part.upper().split(';')[1:]:
        p = part.strip()
        if p in KNOWN_IMAGE_TYPES:
            return p
    return None


def detect_encoding(key_part):
    """
    Return 'QUOTED-PRINTABLE', 'BASE64', or None.
    Handles v3.0+ ENCODING=... and v2.1 inline field;QUOTED-PRINTABLE:.
    """
    match = re.search(r'ENCODING=([A-Za-z0-9-]+)', key_part, re.IGNORECASE)
    if match:
        return match.group(1).upper()
    for part in key_part.upper().split(';')[1:]:
        if part.strip() in ('QUOTED-PRINTABLE', 'BASE64'):
            return part.strip()
    return None


def decode_qp(value):
    """Decode Quoted-Printable, trying UTF-8 then Latin-1."""
    try:
        return quopri.decodestring(value.encode()).decode('utf-8')
    except Exception:
        try:
            return quopri.decodestring(value.encode()).decode('latin-1')
        except Exception:
            return value


def unfold(content):
    """
    Normalize line endings and unfold wrapped vCard lines.
    Per spec, a line starting with a space or tab continues the previous line.
    """
    content = content.replace('\r\n', '\n').replace('\r', '\n')
    result = []
    for line in content.split('\n'):
        if line and line[0] in (' ', '\t'):
            if result:
                result[-1] += line[1:]
        else:
            result.append(line)
    return '\n'.join(result)


def strip_group(prop):
    """Strip vCard group prefix: 'item1.TEL' → 'TEL'."""
    return prop.split('.', 1)[1] if '.' in prop else prop


def unescape(value):
    """Unescape vCard text escape sequences."""
    return (
        value
        .replace('\\n', '\n')
        .replace('\\N', '\n')
        .replace('\\,', ',')
        .replace('\\;', ';')
    )


def format_date(value):
    """
    Format a vCard date value to human-readable.
      YYYYMMDD / YYYY-MM-DD  →  April 12, 1985
      --MMDD / --MM-DD       →  July 4  (no year)
      YYYY                   →  1985
      Unrecognized           →  returned as-is
    """
    if not value:
        return value

    date_part = value.split('T')[0]

    if date_part.startswith('--'):
        digits = date_part[2:].replace('-', '')
        if len(digits) == 4 and digits.isdigit():
            m, d = int(digits[:2]), int(digits[2:])
            if 1 <= m <= 12 and 1 <= d <= 31:
                return f"{MONTHS[m - 1]} {d}"
        return value

    # Handle slash-form dates (M/D/Y or D/M/Y)
    if '/' in date_part:
        parts = date_part.split('/')
        if len(parts) == 3:
            try:
                a, b, c = int(parts[0]), int(parts[1]), int(parts[2])
                # Detect format: if first > 12, must be D/M/Y (day/month/year)
                if a > 12:
                    d, m, y = a, b, c
                # If second > 12, must be M/D/Y with b as day
                elif b > 12:
                    m, d, y = a, b, c
                # Ambiguous: assume US format M/D/Y
                else:
                    m, d, y = a, b, c
                # Handle 2-digit years: 00-30 → 2000-2030, 31-99 → 1931-1999
                if y < 100:
                    y = 2000 + y if y <= 30 else 1900 + y
                if 1 <= m <= 12 and 1 <= d <= 31:
                    return f"{MONTHS[m - 1]} {d}, {y}"
            except (ValueError, IndexError):
                pass
        return value

    # Handle numeric dates (YYYYMMDD, YYYYMM, YYYY)
    digits = date_part.replace('-', '')
    if digits.isdigit():
        if len(digits) == 8:
            y, m, d = int(digits[:4]), int(digits[4:6]), int(digits[6:])
            if 1 <= m <= 12 and 1 <= d <= 31:
                return f"{MONTHS[m - 1]} {d}, {y}"
        elif len(digits) == 6:
            # YYYYMM format (month/year only)
            y, m = int(digits[:4]), int(digits[4:6])
            if 1 <= m <= 12:
                return f"{MONTHS[m - 1]} {y}"
        elif len(digits) == 4:
            return digits

    return value


def format_datetime(value):
    """
    Format a vCard datetime value to human-readable with time.
      YYYYMMDDTHHMMSSz  →  April 12, 1985, 12:00 UTC
      Date only         →  delegates to format_date
    """
    if not value:
        return value
    if 'T' not in value:
        return format_date(value)

    date_part, time_part = value.split('T', 1)
    formatted = format_date(date_part)
    time_digits = re.sub(r'[^0-9]', '', time_part)
    if len(time_digits) >= 4:
        hh, mm = time_digits[:2], time_digits[2:4]
        tz = ' UTC' if time_part.upper().endswith('Z') else ''
        return f"{formatted}, {hh}:{mm}{tz}"
    return formatted


def safe_output_path(path):
    """
    Return a Path that does not yet exist.
    If path exists, appends _1, _2, ... until a free name is found.
    """
    p = Path(path)
    if not p.exists():
        return p
    counter = 1
    while True:
        candidate = p.parent / f"{p.stem}_{counter}{p.suffix}"
        if not candidate.exists():
            return candidate
        counter += 1


def read_file(path, encoding=None):
    """
    Read a file's text content with smart encoding detection.

    If encoding is given (via --encoding flag), it is used directly —
    no detection, no fallback. An invalid encoding or decode failure
    raises an informative error.

    Otherwise, detection runs in this order:
      1. charset-normalizer  — best-in-class detection, used silently if installed
      2. Sequence fallback   — tries each encoding strictly before moving on:
           utf-8-sig          UTF-8 with BOM (Windows tools often add this)
           utf-8              Modern standard; strict, so bad bytes fail fast
           utf-16             Structural BOM makes false positives impossible
           windows-1252       Western European superset of Latin-1
           latin-1            Never fails — absolute safety net
      3. latin-1 + replace   Final safety net: unknown bytes become '?' rather
                             than crashing. Should never be reached in practice.

    Known limitation: CJK encodings (Shift-JIS, GB2312, EUC-KR) without
    charset-normalizer installed will silently decode as windows-1252 garbage.
    Workaround: install charset-normalizer, or use --encoding shift-jis.
    """
    raw = path.read_bytes()

    if encoding:
        try:
            return raw.decode(encoding)
        except (UnicodeDecodeError, LookupError) as e:
            raise ValueError(f"Could not decode {path.name} with encoding '{encoding}': {e}")

    # Optional enhancement — use charset-normalizer if installed
    try:
        from charset_normalizer import from_bytes
        result = from_bytes(raw).best()
        if result and result.encoding:
            return str(result)
    except ImportError:
        pass  # not installed — fall through to sequence silently

    # Sequence fallback — zero dependencies
    _ENCODINGS_TO_TRY = ['utf-8-sig', 'utf-8', 'utf-16', 'windows-1252', 'latin-1']
    for enc in _ENCODINGS_TO_TRY:
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue

    # Absolute safety net — should never reach here given latin-1 never fails
    return raw.decode('latin-1', errors='replace')


def apply_filters(contacts, filter_conditions):
    """
    Filter a list of contacts, keeping only those that match ALL conditions.

    Each condition is a string in the form 'field=value'.

    Supported conditions:
      name=John       — case-insensitive substring match on contact name
      org=Acme        — case-insensitive substring match on any organization
      category=Work   — case-insensitive substring match on categories string
      has=phone       — contact has at least one phone number
      has=email       — contact has at least one email address
      has=birthday    — contact has a birthday
      has=address     — contact has at least one address
      has=note        — contact has at least one note
      has=url         — contact has at least one URL

    Multiple conditions are ANDed — a contact must satisfy all of them.
    Unknown condition keys raise ValueError with a helpful message.
    Uses partition('=') so values containing '=' are handled correctly.

    Returns the filtered list (may be empty).
    """
    VALID_HAS    = {'phone', 'email', 'birthday', 'address', 'note', 'url'}
    VALID_FIELDS = {'name', 'org', 'category', 'has'}

    # Parse and validate all conditions before filtering — fail fast on bad input
    parsed = []
    for cond in filter_conditions:
        if '=' not in cond:
            raise ValueError(
                f"Invalid filter '{cond}'. Use 'field=value' format "
                f"(e.g. 'name=John', 'has=phone', 'org=Acme')."
            )
        key, _, value = cond.partition('=')
        key   = key.strip().lower()
        value = value.strip()
        if not value:
            raise ValueError(f"Filter '{cond}' has no value after '='.")
        if key not in VALID_FIELDS:
            raise ValueError(
                f"Unknown filter field '{key}'. "
                f"Valid fields: {', '.join(sorted(VALID_FIELDS))}."
            )
        if key == 'has' and value.lower() not in VALID_HAS:
            raise ValueError(
                f"Unknown 'has' value '{value}'. "
                f"Valid values: {', '.join(sorted(VALID_HAS))}."
            )
        parsed.append((key, value.lower()))

    result = []
    for contact in contacts:
        match = True
        for key, value in parsed:
            if key == 'name':
                if value not in (contact.get('name') or '').lower():
                    match = False; break
            elif key == 'org':
                orgs = ' '.join(contact.get('organizations') or []).lower()
                if value not in orgs:
                    match = False; break
            elif key == 'category':
                if value not in (contact.get('categories') or '').lower():
                    match = False; break
            elif key == 'has':
                checks = {
                    'phone':    bool(contact.get('phones')),
                    'email':    bool(contact.get('emails')),
                    'birthday': bool(contact.get('birthday')),
                    'address':  bool(contact.get('addresses')),
                    'note':     bool(contact.get('notes')),
                    'url':      bool(contact.get('urls')),
                }
                if not checks.get(value, False):
                    match = False; break
        if match:
            result.append(contact)

    return result


def apply_selection(contacts, select_expr, exclude=False):
    """
    Select or exclude contacts by index, range, last-N, or name wildcard.

    select_expr is a comma-separated string of tokens:
      1           individual 1-based contact number
      1-10        inclusive range
      1-10,15     mix of ranges and individuals
      last-10     last N contacts in the current list
      John*       fnmatch wildcard matched against contact name (case-insensitive)

    Token types can be mixed freely: "1-5,John*,20"

    If exclude=True, the matched contacts are removed instead of kept.
    Order of remaining contacts is always preserved.

    Raises ValueError with a descriptive message if:
      - An index is out of range
      - A range is empty or backwards
      - A token cannot be parsed
    """
    import fnmatch

    total = len(contacts)
    selected_indices = set()  # 0-based internally

    for token in select_expr.split(','):
        token = token.strip()
        if not token:
            continue

        # last-N — must be checked before range to avoid 'last-10' hitting range branch
        if token.lower().startswith('last-'):
            try:
                n = int(token[5:])
                if n < 1:
                    raise ValueError
            except ValueError:
                raise ValueError(
                    f"Invalid last-N token '{token}'. "
                    f"Use last-N where N is a positive integer (e.g. last-10)."
                )
            selected_indices.update(range(max(0, total - n), total))

        # Numeric range: must match digits-dash-digits exactly
        elif re.match(r'^\d+-\d+$', token):
            lo_str, hi_str = token.split('-', 1)
            lo, hi = int(lo_str), int(hi_str)
            if lo < 1:
                raise ValueError(f"Range start must be ≥ 1 (got '{token}').")
            if lo > hi:
                raise ValueError(f"Range '{token}' is backwards. Use '{hi}-{lo}'.")
            if hi > total:
                raise ValueError(
                    f"Contact {hi} is out of range — only {total} "
                    f"{'contact' if total == 1 else 'contacts'} available."
                )
            selected_indices.update(range(lo - 1, hi))

        # Name wildcard: contains * or ?
        elif '*' in token or '?' in token:
            pattern = token.lower()
            for i, c in enumerate(contacts):
                name = (c.get('name') or '').lower()
                if fnmatch.fnmatch(name, pattern):
                    selected_indices.add(i)

        # Single integer index
        else:
            try:
                idx = int(token)
            except ValueError:
                raise ValueError(
                    f"Invalid token '{token}'. "
                    f"Expected an integer, a range (e.g. 1-10), last-N, or a name pattern."
                )
            if idx < 1 or idx > total:
                raise ValueError(
                    f"Contact {idx} is out of range — only {total} "
                    f"{'contact' if total == 1 else 'contacts'} available."
                )
            selected_indices.add(idx - 1)

    if exclude:
        return [c for i, c in enumerate(contacts) if i not in selected_indices]
    else:
        # Preserve original order
        return [c for i, c in enumerate(contacts) if i in selected_indices]


def has_displayable_content(contact):
    """
    Return True if the contact has at least one field worth showing.
    A contact with only _skipped, _version, _source, or empty custom fields
    is treated as having no displayable content and counted as malformed.
    """
    for field in DISPLAYABLE_FIELDS:
        if bool(contact.get(field)):
            return True
    # custom fields: check at least one has a non-empty value
    if any(v for v in contact.get('_custom', {}).values()):
        return True
    return False


# ---------------------------------------------------------------------------
# Parser
# ---------------------------------------------------------------------------

def parse_vcard(content, no_photos=False):
    """
    Parse a single vCard block. Returns a contact dict, or {} if unparseable.
    The vCard version string (e.g. '3.0') is stored internally as '_version'.

    no_photos : if True, PHOTO/LOGO fields are skipped without decoding
                (matches the SILENT_FIELDS skip — no warning, just faster
                on files with many/large embedded images).
    """
    content = unfold(content)
    contact = {}
    skipped = []
    custom = {}
    version = None

    for line in content.strip().split('\n'):
        line = line.strip()
        if not line or ':' not in line:
            continue

        key_part, value = line.split(':', 1)
        key_part = key_part.strip()
        value = value.strip()

        prop = strip_group(key_part.split(';')[0]).upper()

        # Capture version for reporting
        if prop == 'VERSION':
            version = value
            continue

        encoding = detect_encoding(key_part)

        if encoding == 'BASE64':
            if prop in ('PHOTO', 'LOGO'):
                if not no_photos:
                    try:
                        image_bytes = base64.b64decode(value, validate=False)
                        media_type = extract_media_type(key_part) or 'UNKNOWN'
                        if prop == 'PHOTO':
                            contact['_photo_data'] = image_bytes
                            contact['_photo_type'] = media_type
                        else:
                            contact['_logo_data'] = image_bytes
                            contact['_logo_type'] = media_type
                    except Exception as e:
                        skipped.append(f"{prop}: Could not decode image data ({e})")
                # no_photos=True: skip silently, same as a SILENT_FIELDS entry
            elif prop not in SILENT_FIELDS:
                skipped.append(f"{prop}: Base64 encoded data (skipped)")
            continue

        if encoding == 'QUOTED-PRINTABLE':
            value = decode_qp(value)

        if prop in SILENT_FIELDS:
            continue

        try:
            if prop == 'FN':
                contact['name'] = value

            elif prop == 'N':
                if not contact.get('name'):
                    parts = value.split(';')
                    # vCard N field order: family;given;additional;prefix;suffix
                    # We reassemble as: prefix given additional family suffix
                    name_parts = [
                        parts[i] for i in (3, 1, 2, 0, 4)
                        if i < len(parts) and parts[i]
                    ]
                    contact['name'] = ' '.join(name_parts).strip()

            elif prop == 'TEL':
                if not value:
                    raise ValueError("Empty phone number")
                value = strip_uri_prefix(value)
                if sum(c.isdigit() for c in value) < 5:
                    raise ValueError(f"Too short: {value}")
                label = extract_type(key_part)
                entry = f"Phone: {value} ({label})" if label else f"Phone: {value}"
                contact.setdefault('phones', []).append(entry)

            elif prop == 'EMAIL':
                if not value:
                    raise ValueError("Empty email")
                value = strip_uri_prefix(value)
                if value.count('@') != 1:
                    raise ValueError(f"Invalid email: {value}")
                local, domain = value.split('@')
                if not local or not domain or '.' not in domain:
                    raise ValueError(f"Invalid email: {value}")
                label = extract_type(key_part)
                entry = f"Email: {value} ({label})" if label else f"Email: {value}"
                contact.setdefault('emails', []).append(entry)

            elif prop in ('ADR', 'ADDR'):
                parts = [unescape(p) for p in value.split(';')]
                addr_parts = [parts[i] for i in range(2, 7) if i < len(parts) and parts[i]]
                address = ', '.join(addr_parts)
                if address:
                    label = extract_type(key_part)
                    entry = f"{address} ({label})" if label else address
                    contact.setdefault('addresses', []).append(entry)

            elif prop == 'ORG':
                if value:
                    contact.setdefault('organizations', []).append(value)

            elif prop == 'TITLE':
                if value:
                    contact['title'] = value

            elif prop == 'URL':
                if value:
                    contact.setdefault('urls', []).append(value)

            elif prop == 'NOTE':
                if value:
                    contact.setdefault('notes', []).append(unescape(value))

            elif prop == 'BDAY':
                if not value:
                    raise ValueError("Empty birthday")
                clean = value.replace('-', '').split('T')[0]
                valid = (
                    # YYYYMMDD or YYYY-MM-DD — full date
                    (clean.isdigit() and len(clean) == 8) or
                    # YYYYMM or YYYY-MM — month/year only
                    (clean.isdigit() and len(clean) == 6) or
                    # YYYY — year only
                    (clean.isdigit() and len(clean) == 4) or
                    # --MMDD or --MM-DD — vCard 4.0 no-year format
                    (value.startswith('--') and value[2:].replace('-', '').isdigit()) or
                    # Slash-form: M/D/Y or D/M/Y
                    ('/' in value and len(value.split('/')) == 3) or
                    # Month-name forms (e.g. "April 12, 1985")
                    any(m in value.lower() for m in [
                        'jan', 'feb', 'mar', 'apr', 'may', 'jun',
                        'jul', 'aug', 'sep', 'oct', 'nov', 'dec',
                    ])
                )
                if not valid:
                    raise ValueError(f"Unrecognized format: {value}")
                contact['birthday'] = format_date(value)

            elif prop == 'NICKNAME':
                if value:
                    contact['nickname'] = value

            elif prop == 'LABEL':
                if value:
                    contact['label'] = unescape(value)

            elif prop == 'ANNIVERSARY':
                if value:
                    contact['anniversary'] = format_date(value)

            elif prop == 'GENDER':
                if value:
                    contact['gender'] = value

            elif prop == 'CATEGORIES':
                if value:
                    cats = [c.strip() for c in value.split(',') if c.strip()]
                    if cats:
                        contact['categories'] = ', '.join(cats)

            elif prop == 'CREATED':
                if value:
                    contact['created'] = format_datetime(value)
                    contact['_created_raw'] = value  # preserve raw for merge_contacts

            elif prop in ('REVISED', 'REV'):
                if value:
                    contact['revised'] = format_datetime(value)
                    contact['_revised_raw'] = value  # preserve raw for is_newer and merge_contacts

            elif prop == 'TZ':
                if value:
                    contact['timezone'] = value

            elif prop == 'GEO':
                if value:
                    contact['geo'] = value

            elif prop == 'ROLE':
                if value:
                    contact['role'] = value

            elif prop == 'IMPP':
                # Not stripped via strip_uri_prefix — the scheme (xmpp:, skype:,
                # sip:, etc.) identifies which IM protocol this handle is for.
                if value:
                    contact.setdefault('impp', []).append(value)

            else:
                if prop and prop not in STANDARD_FIELDS:
                    clean_prop = prop[2:] if prop.startswith('X-') else prop
                    if clean_prop == 'LABEL':
                        if value:
                            contact['label'] = unescape(value)
                    else:
                        if value:
                            custom[clean_prop.replace('-', ' ').title()] = value

        except (ValueError, IndexError, AttributeError) as e:
            skipped.append(f"{prop}: {e}")

    # Remove duplicate values within each list field
    for key in ('phones', 'emails', 'addresses', 'urls', 'notes', 'organizations', 'impp'):
        if key in contact:
            seen = []
            for item in contact[key]:
                if item not in seen:
                    seen.append(item)
            contact[key] = seen

    if skipped:
        contact['_skipped'] = skipped
    if custom:
        contact['_custom'] = custom
    if version:
        contact['_version'] = version

    return contact


# ---------------------------------------------------------------------------
# Formatter
# ---------------------------------------------------------------------------

def format_contact(contact):
    """Render a contact dict as a human-readable text block."""
    lines = []

    lines.append(f"Name: {contact['name']}" if contact.get('name') else "Name: (Unknown)")

    for org in contact.get('organizations', []):
        lines.append(f"Organization: {org}")

    if contact.get('title'):
        lines.append(f"Title: {contact['title']}")

    for phone in contact.get('phones', []):
        lines.append(phone)

    for email in contact.get('emails', []):
        lines.append(email)

    for addr in contact.get('addresses', []):
        lines.append(f"Address: {addr}")

    for url in contact.get('urls', []):
        lines.append(f"Website: {url}")

    if contact.get('birthday'):
        lines.append(f"Birthday: {contact['birthday']}")

    for note in contact.get('notes', []):
        lines.append(f"Note: {note}")

    if contact.get('nickname'):
        lines.append(f"Nickname: {contact['nickname']}")

    if contact.get('label'):
        lines.append(f"Label: {contact['label']}")

    if contact.get('anniversary'):
        lines.append(f"Anniversary: {contact['anniversary']}")

    if contact.get('gender'):
        lines.append(f"Gender: {contact['gender']}")

    if contact.get('categories'):
        lines.append(f"Categories: {contact['categories']}")

    if contact.get('created'):
        lines.append(f"Created: {contact['created']}")

    if contact.get('revised'):
        lines.append(f"Revised: {contact['revised']}")

    for field_name, field_value in contact.get('_custom', {}).items():
        lines.append(f"* {field_name}: {field_value}")

    if contact.get('_skipped'):
        lines.append("")
        lines.append("[Skipped Fields]")
        for msg in contact['_skipped']:
            lines.append(f"  ⚠ {msg}")

    return '\n'.join(lines)


def format_vcf(contact):
    """
    Render a contact dict as a valid vCard 3.0 block.

    Phones and emails are stored as presentation strings (D1 design decision),
    so we reverse-parse them to extract the raw value and type label:
      "Phone: +1-555-0101 (Work)"  →  TEL;TYPE=WORK:+1-555-0101
      "Email: j@x.com (Work)"      →  EMAIL;TYPE=WORK:j@x.com

    _revised_raw / _created_raw are used for REV/CREATED so the output
    contains the original ISO timestamps, not the formatted display strings.

    Custom fields have their display name reversed back to X- form:
      {'Skype': 'john'} → X-SKYPE:john
    """
    lines = ['BEGIN:VCARD', 'VERSION:3.0']

    if contact.get('name'):
        lines.append(f"FN:{contact['name']}")

    for org in contact.get('organizations', []):
        lines.append(f"ORG:{org}")

    if contact.get('title'):
        lines.append(f"TITLE:{contact['title']}")

    # Reverse-parse "Phone: +1-555-0101 (Work/Voice)" → TEL;TYPE=WORK,VOICE:+1-555-0101
    for entry in contact.get('phones', []):
        # Format: "Phone: <value>" or "Phone: <value> (<type>)"
        rest = entry[len('Phone: '):]
        if rest.endswith(')') and ' (' in rest:
            value, type_str = rest.rsplit(' (', 1)
            type_str = type_str[:-1]  # strip trailing )
            type_param = ','.join(p.upper() for p in type_str.split('/'))
            lines.append(f"TEL;TYPE={type_param}:{value}")
        else:
            lines.append(f"TEL:{rest}")

    # Reverse-parse "Email: j@x.com (Work/Internet)" → EMAIL;TYPE=WORK,INTERNET:j@x.com
    for entry in contact.get('emails', []):
        rest = entry[len('Email: '):]
        if rest.endswith(')') and ' (' in rest:
            value, type_str = rest.rsplit(' (', 1)
            type_str = type_str[:-1]
            type_param = ','.join(p.upper() for p in type_str.split('/'))
            lines.append(f"EMAIL;TYPE={type_param}:{value}")
        else:
            lines.append(f"EMAIL:{rest}")

    # Addresses stored as "street, city, state, zip, country (Type)"
    # We write them to ADR field 3 (street) only — full structured reassembly
    # would require storing parts separately (v5 refactor). This is readable
    # and importable; contact apps handle free-form ADR gracefully.
    for entry in contact.get('addresses', []):
        if entry.endswith(')') and ' (' in entry:
            value, type_str = entry.rsplit(' (', 1)
            type_str = type_str[:-1]
            type_param = ','.join(p.upper() for p in type_str.split('/'))
            lines.append(f"ADR;TYPE={type_param}:;;{value};;;;")
        else:
            lines.append(f"ADR:;;{entry};;;;")

    for url in contact.get('urls', []):
        lines.append(f"URL:{url}")

    if contact.get('birthday'):
        lines.append(f"BDAY:{contact['birthday']}")

    for note in contact.get('notes', []):
        # Re-escape newlines for vCard
        escaped = note.replace('\n', '\\n')
        lines.append(f"NOTE:{escaped}")

    if contact.get('nickname'):
        lines.append(f"NICKNAME:{contact['nickname']}")

    if contact.get('label'):
        escaped = contact['label'].replace('\n', '\\n')
        lines.append(f"LABEL:{escaped}")

    if contact.get('anniversary'):
        lines.append(f"ANNIVERSARY:{contact['anniversary']}")

    if contact.get('gender'):
        lines.append(f"GENDER:{contact['gender']}")

    if contact.get('categories'):
        lines.append(f"CATEGORIES:{contact['categories']}")

    # Use raw ISO timestamps for REV/CREATED — not the formatted display strings
    if contact.get('_created_raw'):
        lines.append(f"CREATED:{contact['_created_raw']}")
    elif contact.get('created'):
        lines.append(f"CREATED:{contact['created']}")

    if contact.get('_revised_raw'):
        lines.append(f"REV:{contact['_revised_raw']}")
    elif contact.get('revised'):
        lines.append(f"REV:{contact['revised']}")

    # Custom fields: reverse display name back to X- form
    # "Skype" → X-SKYPE, "Twitter Handle" → X-TWITTER-HANDLE
    for field_name, field_value in contact.get('_custom', {}).items():
        prop_name = 'X-' + field_name.upper().replace(' ', '-')
        lines.append(f"{prop_name}:{field_value}")

    lines.append('END:VCARD')
    return '\n'.join(lines)


def contact_to_json_dict(contact, include_photos=True):
    """
    Convert an internal contact dict into a JSON-serializable dict.

    Note (D1 design decision, unchanged in v4): phones and emails are still
    stored as presentation strings ("Phone: +1-555-0101 (Work)") rather than
    structured {'value':..., 'type':...} dicts — that refactor is deferred
    to the v5 D1 rework so parser/formatter/dedup change together. JSON
    consumers get the same presentation strings as the text output for now.

    include_photos : if False, PHOTO/LOGO data is omitted even if the
                      contact has it (mirrors --no-photos at output time,
                      independent of whether parse_vcard() decoded it).

    Keys with empty/None values are omitted to keep the output clean.
    """
    d = {
        'name':          contact.get('name'),
        'organizations': contact.get('organizations', []),
        'title':         contact.get('title'),
        'phones':        contact.get('phones', []),
        'emails':        contact.get('emails', []),
        'addresses':     contact.get('addresses', []),
        'urls':          contact.get('urls', []),
        'birthday':      contact.get('birthday'),
        'notes':         contact.get('notes', []),
        'nickname':      contact.get('nickname'),
        'label':         contact.get('label'),
        'anniversary':   contact.get('anniversary'),
        'gender':        contact.get('gender'),
        'categories':    contact.get('categories'),
        'created':       contact.get('created'),
        'revised':       contact.get('revised'),
        'role':          contact.get('role'),
        'timezone':      contact.get('timezone'),
        'geo':           contact.get('geo'),
        'impp':          contact.get('impp', []),
        'custom_fields': contact.get('_custom', {}),
        'vcard_version': contact.get('_version'),
        'source_file':   contact.get('_source'),
        'skipped_fields': contact.get('_skipped', []),
    }

    if include_photos:
        if contact.get('_photo_data') is not None:
            d['photo'] = {
                'type':        contact.get('_photo_type'),
                'data_base64': base64.b64encode(contact['_photo_data']).decode('ascii'),
            }
        if contact.get('_logo_data') is not None:
            d['logo'] = {
                'type':        contact.get('_logo_type'),
                'data_base64': base64.b64encode(contact['_logo_data']).decode('ascii'),
            }

    return {k: v for k, v in d.items() if v not in (None, '', [], {})}


def format_json(all_contacts, meta, include_photos=True):
    """
    Render the full run as a JSON document: a metadata envelope plus the
    contact array. Matches the informational depth of the text/VCF output
    headers (see convert()'s header_rows) but as structured key-value pairs
    instead of display-formatted strings.
    """
    payload = {
        'meta': meta,
        'contacts': [contact_to_json_dict(c, include_photos=include_photos) for c in all_contacts],
    }
    return json.dumps(payload, indent=2, ensure_ascii=False)


def _render_text(all_contacts, header_rows, single_mode, malformed, exact_removed, fuzzy_warnings):
    """
    Render the v2/v3-style human-readable text output. Extracted verbatim
    from convert()'s original inline text-building code so the v4 format
    dispatcher can call it the same way it calls every other formatter —
    behavior and output are unchanged from v3.
    """
    col_width = max(len(r[0]) for r in header_rows)
    header_lines = [f"  {label:<{col_width}}  :  {value}" for label, value in header_rows]

    lines = []
    lines.append("Generated by vCardTra")
    lines.append(DIVIDER)
    lines.extend(header_lines)
    lines.append(DIVIDER)
    lines.append("")

    for i, contact in enumerate(all_contacts, 1):
        if single_mode:
            header = f"⭐ Contact {i}:"
        else:
            header = f"⭐ Contact {i}: (from {contact.get('_source', '?')})"
        lines.append(header)
        lines.append(DIVIDER_CTX)
        lines.append(format_contact(contact))
        lines.append("")

    warn_lines = []
    if malformed:
        warn_lines.append(f"  ⚠ {malformed} malformed vCard block(s) skipped (no parseable content)")
    if exact_removed:
        warn_lines.append(f"  ⚠ {len(exact_removed)} exact duplicate(s) removed (newer revision kept where available):")
        for retained_idx, name in exact_removed:
            warn_lines.append(f"      - {name} → merged into Contact {retained_idx}")
    if fuzzy_warnings:
        for i, j, name in fuzzy_warnings:
            warn_lines.append(f"  ⚠ Duplicate warning: Contact {i} and Contact {j} ({name})")

    if warn_lines:
        lines.append(DIVIDER_HVY)
        lines.append("[Conversion Warnings]")
        lines.extend(warn_lines)
        lines.append(DIVIDER_HVY)

    return '\n'.join(lines)


def _strip_field_prefix(value, prefix):
    """
    Strip a leading 'Phone: ' / 'Email: ' style prefix from a presentation
    string, for formats (CSV, SQL, XML) where the column/tag name already
    says what the field is — text/Markdown/HTML keep the prefix since it
    reads naturally inline there.
    """
    return value[len(prefix):] if value.startswith(prefix) else value


def contact_to_flat_dict(contact):
    """
    Convert a contact dict into a flat, single-valued dict for row-based
    formats (CSV wide columns, SQL/SQLite single table, XML attributes).
    Multi-value fields (organizations, notes, impp) are joined with ' | ';
    phones/emails/addresses/urls are handled separately by callers that
    need numbered columns (CSV) since a flat join loses the per-value
    structure those formats want to preserve.

    Same D1 caveat as contact_to_json_dict(): phones/emails stay as
    presentation strings, not structured dicts (deferred to v5).
    """
    return {
        'name':            contact.get('name') or '',
        'organizations':   ' | '.join(contact.get('organizations', [])),
        'title':           contact.get('title') or '',
        'birthday':        contact.get('birthday') or '',
        'notes':           ' | '.join(contact.get('notes', [])),
        'nickname':        contact.get('nickname') or '',
        'label':           contact.get('label') or '',
        'anniversary':     contact.get('anniversary') or '',
        'gender':          contact.get('gender') or '',
        'categories':      contact.get('categories') or '',
        'created':         contact.get('created') or '',
        'revised':         contact.get('revised') or '',
        'role':            contact.get('role') or '',
        'timezone':        contact.get('timezone') or '',
        'geo':             contact.get('geo') or '',
        'impp':            ' | '.join(contact.get('impp', [])),
        'custom_fields':   ' | '.join(f"{k}: {v}" for k, v in contact.get('_custom', {}).items()),
        'vcard_version':   contact.get('_version') or '',
        'source_file':     contact.get('_source') or '',
        'skipped_fields':  ' | '.join(contact.get('_skipped', [])),
    }


def format_csv(all_contacts):
    """
    Render contacts as a wide CSV: one row per contact, numbered columns
    for multi-value fields (Phone1, Phone2, ... Email1, Email2, ...) since
    CSV has no native concept of a repeated field. Column count is set by
    the contact with the most values for that field across the whole run.

    Photos/logos are skipped entirely (binary data has no sane CSV
    representation) — matches the v4 plan table.
    """
    max_phones    = max((len(c.get('phones', []))    for c in all_contacts), default=0)
    max_emails    = max((len(c.get('emails', []))     for c in all_contacts), default=0)
    max_addresses = max((len(c.get('addresses', []))  for c in all_contacts), default=0)
    max_urls      = max((len(c.get('urls', []))       for c in all_contacts), default=0)

    fieldnames = ['Name', 'Organizations', 'Title']
    fieldnames += [f'Phone{i+1}' for i in range(max_phones)]
    fieldnames += [f'Email{i+1}' for i in range(max_emails)]
    fieldnames += [f'Address{i+1}' for i in range(max_addresses)]
    fieldnames += ['Birthday']
    fieldnames += [f'URL{i+1}' for i in range(max_urls)]
    fieldnames += ['Notes', 'Nickname', 'Label', 'Anniversary', 'Gender', 'Categories',
                   'Created', 'Revised', 'Role', 'Timezone', 'Geo', 'IMPP',
                   'Custom Fields', 'VCard Version', 'Source File', 'Skipped Fields']

    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=fieldnames, extrasaction='ignore')
    writer.writeheader()

    for c in all_contacts:
        row = contact_to_flat_dict(c)
        row['Name'] = row.pop('name')
        row['Organizations'] = row.pop('organizations')
        row['Title'] = row.pop('title')
        row['Birthday'] = row.pop('birthday')
        row['Notes'] = row.pop('notes')
        row['Nickname'] = row.pop('nickname')
        row['Label'] = row.pop('label')
        row['Anniversary'] = row.pop('anniversary')
        row['Gender'] = row.pop('gender')
        row['Categories'] = row.pop('categories')
        row['Created'] = row.pop('created')
        row['Revised'] = row.pop('revised')
        row['Role'] = row.pop('role')
        row['Timezone'] = row.pop('timezone')
        row['Geo'] = row.pop('geo')
        row['IMPP'] = row.pop('impp')
        row['Custom Fields'] = row.pop('custom_fields')
        row['VCard Version'] = row.pop('vcard_version')
        row['Source File'] = row.pop('source_file')
        row['Skipped Fields'] = row.pop('skipped_fields')

        for i, p in enumerate(c.get('phones', [])):
            row[f'Phone{i+1}'] = _strip_field_prefix(p, 'Phone: ')
        for i, e in enumerate(c.get('emails', [])):
            row[f'Email{i+1}'] = _strip_field_prefix(e, 'Email: ')
        for i, a in enumerate(c.get('addresses', [])):
            row[f'Address{i+1}'] = a
        for i, u in enumerate(c.get('urls', [])):
            row[f'URL{i+1}'] = u

        writer.writerow(row)

    return output.getvalue()


def format_markdown(all_contacts, header_rows, single_mode):
    """
    Render contacts as a single Markdown document — one H2 section per
    contact, fields as a bullet list. Designed for Notion/Obsidian/GitHub
    wikis: headers give a document outline, plain bullets stay readable
    as raw text too. Field order and content mirror format_contact() (text)
    rather than introducing a third field ordering to keep track of.
    """
    lines = ['# vCardTra Export', '']
    for label, value in header_rows:
        lines.append(f"**{label}:** {value}  ")
    lines.append('')
    lines.append('---')
    lines.append('')

    for i, c in enumerate(all_contacts, 1):
        name = c.get('name') or '(Unknown)'
        lines.append(f"## {i}. {name}")
        if not single_mode:
            lines.append(f"*Source: {c.get('_source', '?')}*")
        lines.append('')

        for org in c.get('organizations', []):
            lines.append(f"- **Organization:** {org}")
        if c.get('title'):
            lines.append(f"- **Title:** {c['title']}")
        for phone in c.get('phones', []):
            lines.append(f"- {phone}")
        for email in c.get('emails', []):
            lines.append(f"- {email}")
        for addr in c.get('addresses', []):
            lines.append(f"- **Address:** {addr}")
        for url in c.get('urls', []):
            lines.append(f"- **Website:** {url}")
        if c.get('birthday'):
            lines.append(f"- **Birthday:** {c['birthday']}")
        for note in c.get('notes', []):
            lines.append(f"- **Note:** {note}")
        if c.get('nickname'):
            lines.append(f"- **Nickname:** {c['nickname']}")
        if c.get('label'):
            lines.append(f"- **Label:** {c['label']}")
        if c.get('anniversary'):
            lines.append(f"- **Anniversary:** {c['anniversary']}")
        if c.get('gender'):
            lines.append(f"- **Gender:** {c['gender']}")
        if c.get('categories'):
            lines.append(f"- **Categories:** {c['categories']}")
        if c.get('created'):
            lines.append(f"- **Created:** {c['created']}")
        if c.get('revised'):
            lines.append(f"- **Revised:** {c['revised']}")
        for field_name, field_value in c.get('_custom', {}).items():
            lines.append(f"- *{field_name}:* {field_value}")

        if c.get('_skipped'):
            lines.append('')
            lines.append('**Skipped fields:**')
            for msg in c['_skipped']:
                lines.append(f"- ⚠ {msg}")

        lines.append('')
        lines.append('---')
        lines.append('')

    return '\n'.join(lines)


def format_html(all_contacts, header_rows, single_mode, include_photos=True):
    """
    Render a single self-contained HTML file — no external CSS/JS/images,
    everything (including photos, if present and included) inline so the
    file works standalone when opened in a browser or emailed. Also the
    foundation the v4 plan calls for PDF conversion to build on.
    """
    def esc(s):
        return html.escape(str(s), quote=False)

    parts = []
    parts.append('<!DOCTYPE html>')
    parts.append('<html lang="en"><head><meta charset="UTF-8">')
    parts.append('<title>vCardTra Export</title>')
    parts.append('<style>')
    parts.append('''
body { font-family: -apple-system, Segoe UI, Roboto, sans-serif; max-width: 900px;
       margin: 2rem auto; padding: 0 1rem; color: #1a1a1a; background: #fafafa; }
h1 { font-size: 1.5rem; }
.header-box { background: #fff; border: 1px solid #ddd; border-radius: 8px;
              padding: 1rem 1.25rem; margin-bottom: 1.5rem; }
.header-box div { font-size: 0.9rem; color: #444; padding: 2px 0; }
.header-box b { color: #111; }
.contact { background: #fff; border: 1px solid #ddd; border-radius: 8px;
           padding: 1.25rem 1.5rem; margin-bottom: 1.25rem; display: flex; gap: 1.25rem; }
.contact-photo { width: 96px; height: 96px; border-radius: 8px; object-fit: cover; flex-shrink: 0; }
.contact-body { flex: 1; min-width: 0; }
.contact h2 { margin: 0 0 0.5rem 0; font-size: 1.15rem; }
.contact .source { color: #888; font-size: 0.8rem; margin-bottom: 0.5rem; }
.field { font-size: 0.92rem; padding: 2px 0; }
.field b { color: #333; }
.skipped { margin-top: 0.75rem; font-size: 0.85rem; color: #a15c00; }
.skipped li { margin-left: 1rem; }
''')
    parts.append('</style></head><body>')
    parts.append('<h1>vCardTra Export</h1>')
    parts.append('<div class="header-box">')
    for label, value in header_rows:
        parts.append(f'<div><b>{esc(label)}:</b> {esc(value)}</div>')
    parts.append('</div>')

    for i, c in enumerate(all_contacts, 1):
        name = c.get('name') or '(Unknown)'
        parts.append('<div class="contact">')

        photo_bytes = c.get('_photo_data') if include_photos else None
        if photo_bytes:
            mime = f"image/{(c.get('_photo_type') or 'jpeg').lower()}"
            b64 = base64.b64encode(photo_bytes).decode('ascii')
            parts.append(f'<img class="contact-photo" src="data:{mime};base64,{b64}" alt="">')

        parts.append('<div class="contact-body">')
        parts.append(f'<h2>{i}. {esc(name)}</h2>')
        if not single_mode:
            parts.append(f'<div class="source">Source: {esc(c.get("_source", "?"))}</div>')

        for org in c.get('organizations', []):
            parts.append(f'<div class="field"><b>Organization:</b> {esc(org)}</div>')
        if c.get('title'):
            parts.append(f'<div class="field"><b>Title:</b> {esc(c["title"])}</div>')
        if c.get('role'):
            parts.append(f'<div class="field"><b>Role:</b> {esc(c["role"])}</div>')
        for phone in c.get('phones', []):
            parts.append(f'<div class="field">{esc(phone)}</div>')
        for email in c.get('emails', []):
            parts.append(f'<div class="field">{esc(email)}</div>')
        for addr in c.get('addresses', []):
            parts.append(f'<div class="field"><b>Address:</b> {esc(addr)}</div>')
        for url in c.get('urls', []):
            parts.append(f'<div class="field"><b>Website:</b> <a href="{esc(url)}">{esc(url)}</a></div>')
        if c.get('geo'):
            lat_lon = c['geo'].replace(';', ',')
            maps_url = f"https://www.google.com/maps?q={lat_lon}"
            parts.append(f'<div class="field"><b>Location:</b> <a href="{esc(maps_url)}">{esc(c["geo"])}</a></div>')
        if c.get('timezone'):
            parts.append(f'<div class="field"><b>Timezone:</b> {esc(c["timezone"])}</div>')
        if c.get('birthday'):
            parts.append(f'<div class="field"><b>Birthday:</b> {esc(c["birthday"])}</div>')
        for handle in c.get('impp', []):
            parts.append(f'<div class="field"><b>IM:</b> {esc(handle)}</div>')
        for note in c.get('notes', []):
            parts.append(f'<div class="field"><b>Note:</b> {esc(note)}</div>')
        if c.get('nickname'):
            parts.append(f'<div class="field"><b>Nickname:</b> {esc(c["nickname"])}</div>')
        if c.get('anniversary'):
            parts.append(f'<div class="field"><b>Anniversary:</b> {esc(c["anniversary"])}</div>')
        if c.get('gender'):
            parts.append(f'<div class="field"><b>Gender:</b> {esc(c["gender"])}</div>')
        if c.get('categories'):
            parts.append(f'<div class="field"><b>Categories:</b> {esc(c["categories"])}</div>')
        for field_name, field_value in c.get('_custom', {}).items():
            parts.append(f'<div class="field"><i>{esc(field_name)}:</i> {esc(field_value)}</div>')

        if c.get('_skipped'):
            parts.append('<ul class="skipped">')
            for msg in c['_skipped']:
                parts.append(f'<li>⚠ {esc(msg)}</li>')
            parts.append('</ul>')

        parts.append('</div></div>')

    parts.append('</body></html>')
    return '\n'.join(parts)


def parse_display_date(value):
    """
    Reverse-parse format_date()'s human-readable output back into
    (year, month, day) — any of which may be None. Needed by the ICS
    formatter, which requires a real calendar date, but birthday/anniversary
    are stored post-formatting (see Design Decisions: 'format_date called
    at parse time'). Only handles the exact output shapes format_date()
    produces; anything else (unparsed passthrough values) returns all-None.

    Recognised shapes:
      "April 12, 1985"  -> (1985, 4, 12)
      "July 4"          -> (None, 7, 4)      -- from --MMDD, no year
      "April 1985"      -> (1985, 4, None)   -- from YYYYMM, no day
      "1985"             -> (1985, None, None)
    """
    if not value:
        return (None, None, None)

    month_lookup = {m.lower(): i + 1 for i, m in enumerate(MONTHS)}

    # "Month Day, Year"
    m = re.match(r'^([A-Za-z]+)\s+(\d{1,2}),\s+(\d{4})$', value)
    if m and m.group(1).lower() in month_lookup:
        return (int(m.group(3)), month_lookup[m.group(1).lower()], int(m.group(2)))

    # "Month Day" (no year)
    m = re.match(r'^([A-Za-z]+)\s+(\d{1,2})$', value)
    if m and m.group(1).lower() in month_lookup:
        return (None, month_lookup[m.group(1).lower()], int(m.group(2)))

    # "Month Year" (no day)
    m = re.match(r'^([A-Za-z]+)\s+(\d{4})$', value)
    if m and m.group(1).lower() in month_lookup:
        return (int(m.group(2)), month_lookup[m.group(1).lower()], None)

    # Bare year
    if re.match(r'^\d{4}$', value):
        return (int(value), None, None)

    return (None, None, None)


def format_ics(all_contacts):
    """
    Render birthdays and anniversaries as a yearly-recurring ICS calendar.
    Only dates with both a month and a day can become a calendar event —
    month-only (YYYYMM) or year-only birthdays have no day to put the event
    on, so those are skipped (not an error, just not calendar-representable).

    Events with a known year use it as DTSTART's year (informational only,
    since RRULE:FREQ=YEARLY repeats regardless); events with no year use
    1900 as a conventional placeholder, following the common iCal practice
    for "unknown year" recurring reminders, and say so in the description.
    """
    lines = ['BEGIN:VCALENDAR', 'VERSION:2.0', 'PRODID:-//vCardTra//v4//EN', 'CALSCALE:GREGORIAN']

    def add_event(name, kind, year, month, day, uid_suffix):
        display_year = year if year else 1900
        dtstart = f"{display_year:04d}{month:02d}{day:02d}"
        summary = f"{name} — {kind}" + ('' if year else ' (year unknown)')
        lines.append('BEGIN:VEVENT')
        lines.append(f"UID:{uid_suffix}@vcardtra")
        lines.append(f"DTSTART;VALUE=DATE:{dtstart}")
        lines.append(f"DTEND;VALUE=DATE:{dtstart}")
        lines.append('RRULE:FREQ=YEARLY')
        lines.append(f"SUMMARY:{_ics_escape(summary)}")
        lines.append('END:VEVENT')

    skipped = 0
    for i, c in enumerate(all_contacts, 1):
        name = c.get('name') or '(Unknown)'
        if c.get('birthday'):
            year, month, day = parse_display_date(c['birthday'])
            if month and day:
                add_event(name, 'Birthday', year, month, day, f"bday-{i}")
            else:
                skipped += 1
        if c.get('anniversary'):
            year, month, day = parse_display_date(c['anniversary'])
            if month and day:
                add_event(name, 'Anniversary', year, month, day, f"anniv-{i}")
            else:
                skipped += 1

    lines.append('END:VCALENDAR')
    return '\n'.join(lines), skipped


def _ics_escape(value):
    """Escape text per RFC 5545 — comma, semicolon, backslash, newline."""
    return (
        value.replace('\\', '\\\\')
             .replace(',', '\\,')
             .replace(';', '\\;')
             .replace('\n', '\\n')
    )


def _sql_escape(value):
    """Escape a value for a single-quoted SQL string literal."""
    return str(value).replace("'", "''")


SQL_COLUMNS = [
    'id', 'name', 'organizations', 'title', 'phones', 'emails', 'addresses',
    'birthday', 'urls', 'notes', 'nickname', 'label', 'anniversary', 'gender',
    'categories', 'created', 'revised', 'role', 'timezone', 'geo', 'impp',
    'custom_fields', 'vcard_version', 'source_file', 'skipped_fields',
]


def _contact_sql_row(idx, contact):
    """Build one row's values in SQL_COLUMNS order, for SQL and SQLite."""
    flat = contact_to_flat_dict(contact)
    row = [idx, flat['name'], flat['organizations'], flat['title']]
    row.append(' | '.join(_strip_field_prefix(p, 'Phone: ') for p in contact.get('phones', [])))
    row.append(' | '.join(_strip_field_prefix(e, 'Email: ') for e in contact.get('emails', [])))
    row.append(' | '.join(contact.get('addresses', [])))
    row.append(flat['birthday'])
    row.append(' | '.join(contact.get('urls', [])))
    row += [flat['notes'], flat['nickname'], flat['label'], flat['anniversary'], flat['gender'],
            flat['categories'], flat['created'], flat['revised'], flat['role'], flat['timezone'],
            flat['geo'], flat['impp'], flat['custom_fields'], flat['vcard_version'],
            flat['source_file'], flat['skipped_fields']]
    return row


def format_sql(all_contacts):
    """
    Render a single wide 'contacts' table as plain .sql text: one
    CREATE TABLE plus one INSERT per contact. Same single-table, D1-deferred
    shape as CSV/XML (see contact_to_flat_dict) — normalizing phones/emails
    into their own tables is a v5 D1-refactor-scale change, not a v4 one.
    """
    lines = ['-- Generated by vCardTra', 'BEGIN TRANSACTION;', '']
    lines.append('CREATE TABLE IF NOT EXISTS contacts (')
    col_defs = [f"    {col} {'INTEGER PRIMARY KEY' if col == 'id' else 'TEXT'}" for col in SQL_COLUMNS]
    lines.append(',\n'.join(col_defs))
    lines.append(');')
    lines.append('')

    for idx, c in enumerate(all_contacts, 1):
        row = _contact_sql_row(idx, c)
        values = ', '.join(
            str(v) if isinstance(v, int) else f"'{_sql_escape(v)}'" for v in row
        )
        lines.append(f"INSERT INTO contacts ({', '.join(SQL_COLUMNS)}) VALUES ({values});")

    lines.append('')
    lines.append('COMMIT;')
    return '\n'.join(lines)


def write_sqlite(all_contacts, out_path):
    """
    Write contacts directly to a SQLite .db file via the stdlib sqlite3
    module — same schema as format_sql(), just executed instead of printed
    as text. No server, no extra dependency; sqlite3 ships with Python.
    """
    conn = sqlite3.connect(str(out_path))
    try:
        cur = conn.cursor()
        col_defs = ', '.join(f"{col} {'INTEGER PRIMARY KEY' if col == 'id' else 'TEXT'}" for col in SQL_COLUMNS)
        cur.execute(f"CREATE TABLE IF NOT EXISTS contacts ({col_defs})")
        placeholders = ', '.join('?' for _ in SQL_COLUMNS)
        for idx, c in enumerate(all_contacts, 1):
            row = _contact_sql_row(idx, c)
            cur.execute(f"INSERT INTO contacts ({', '.join(SQL_COLUMNS)}) VALUES ({placeholders})", row)
        conn.commit()
    finally:
        conn.close()


def format_xml(all_contacts, meta, include_photos=True):
    """
    Render the same meta+contacts shape as JSON, as XML. Built with manual
    string assembly (stdlib xml.sax.saxutils.escape for safety) rather than
    ElementTree, since ElementTree's pretty-printing story is awkward before
    Python 3.9's indent() and manual assembly gives more control over the
    output shape without a version-dependent code path.
    """
    def esc(v):
        return xml_escape(str(v))

    def dict_to_xml(d, indent='    '):
        out = []
        for k, v in d.items():
            tag = re.sub(r'[^A-Za-z0-9_]', '_', str(k))
            if isinstance(v, list):
                out.append(f'{indent}<{tag}>')
                for item in v:
                    if isinstance(item, dict):
                        out.append(f'{indent}  <item>')
                        out.append(dict_to_xml(item, indent + '    '))
                        out.append(f'{indent}  </item>')
                    else:
                        out.append(f'{indent}  <item>{esc(item)}</item>')
                out.append(f'{indent}</{tag}>')
            elif isinstance(v, dict):
                out.append(f'{indent}<{tag}>')
                out.append(dict_to_xml(v, indent + '    '))
                out.append(f'{indent}</{tag}>')
            else:
                out.append(f'{indent}<{tag}>{esc(v)}</{tag}>')
        return '\n'.join(out)

    parts = ['<?xml version="1.0" encoding="UTF-8"?>', '<vcardtra_export>', '  <meta>']
    parts.append(dict_to_xml(meta, indent='    '))
    parts.append('  </meta>')
    parts.append('  <contacts>')
    for c in all_contacts:
        d = contact_to_json_dict(c, include_photos=include_photos)
        parts.append('    <contact>')
        parts.append(dict_to_xml(d, indent='      '))
        parts.append('    </contact>')
    parts.append('  </contacts>')
    parts.append('</vcardtra_export>')
    return '\n'.join(parts)


def _validate_image_bytes(data):
    """
    Return True if `data` is fully-decodable image bytes safe to embed.
    Used only by write_pdf() — weasyprint (which depends on Pillow) can raise
    on plausible-but-corrupt image bytes (valid header, truncated/corrupted
    data stream), which would otherwise crash the *entire* multi-contact PDF
    render over one bad photo. Lazy-imports Pillow — already a weasyprint
    dependency, so this adds no new requirement beyond choosing --format pdf;
    if Pillow is somehow unavailable, validation is skipped (same crash risk
    as before, just not made worse).
    """
    try:
        from PIL import Image
        import io as _io
        img = Image.open(_io.BytesIO(data))
        img.load()  # forces full decode — catches truncated/corrupt streams
        return True
    except ImportError:
        return True  # can't validate; don't block rendering over it
    except Exception:
        return False


def write_pdf(all_contacts, header_rows, single_mode, out_path, include_photos=True):
    """
    Render to PDF by building the same HTML as format_html() and letting
    weasyprint lay it out — "free once HTML is done" per the v4 plan.
    weasyprint is optional: not imported at module load, only here, so the
    script still runs with zero dependencies if PDF is never requested.
    Raises ImportError with an install hint if weasyprint isn't present;
    the caller (convert()) catches this the same way as any other write error.
    """
    try:
        from weasyprint import HTML
    except ImportError:
        raise ImportError(
            "--format pdf requires the 'weasyprint' package, which isn't installed.\n"
            "  Install with: pip install weasyprint"
        )

    # Corrupted images get dropped for rendering only (not mutated in the
    # original contact dicts) — see _validate_image_bytes() docstring.
    render_contacts = all_contacts
    if include_photos:
        render_contacts = []
        for c in all_contacts:
            bad_keys = [k for k in ('_photo_data', '_logo_data') if c.get(k) and not _validate_image_bytes(c[k])]
            if bad_keys:
                c = dict(c)
                for k in bad_keys:
                    c[k] = None
            render_contacts.append(c)

    html_content = format_html(render_contacts, header_rows, single_mode, include_photos=include_photos)
    HTML(string=html_content).write_pdf(str(out_path))


def write_docx(all_contacts, header_rows, single_mode, out_path, include_photos=True):
    """
    Render to a Word document via python-docx — one heading + field list
    per contact, photo inserted inline as an image if present. Optional
    dependency, same lazy-import pattern as write_pdf() above.
    """
    try:
        import docx
        from docx.shared import Inches
        import io as _io
    except ImportError:
        raise ImportError(
            "--format docx requires the 'python-docx' package, which isn't installed.\n"
            "  Install with: pip install python-docx"
        )

    document = docx.Document()
    document.add_heading('vCardTra Export', level=1)
    for label, value in header_rows:
        p = document.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(value))

    for i, c in enumerate(all_contacts, 1):
        name = c.get('name') or '(Unknown)'
        document.add_heading(f"{i}. {name}", level=2)
        if not single_mode:
            document.add_paragraph(f"Source: {c.get('_source', '?')}", style='Intense Quote')

        photo_bytes = c.get('_photo_data') if include_photos else None
        if photo_bytes:
            try:
                document.add_picture(_io.BytesIO(photo_bytes), width=Inches(1.2))
            except Exception:
                pass  # unrecognized/corrupt image data — skip the picture, keep the rest

        for org in c.get('organizations', []):
            document.add_paragraph(f"Organization: {org}")
        if c.get('title'):
            document.add_paragraph(f"Title: {c['title']}")
        if c.get('role'):
            document.add_paragraph(f"Role: {c['role']}")
        for phone in c.get('phones', []):
            document.add_paragraph(phone)
        for email in c.get('emails', []):
            document.add_paragraph(email)
        for addr in c.get('addresses', []):
            document.add_paragraph(f"Address: {addr}")
        for url in c.get('urls', []):
            document.add_paragraph(f"Website: {url}")
        if c.get('birthday'):
            document.add_paragraph(f"Birthday: {c['birthday']}")
        for handle in c.get('impp', []):
            document.add_paragraph(f"IM: {handle}")
        for note in c.get('notes', []):
            document.add_paragraph(f"Note: {note}")
        if c.get('nickname'):
            document.add_paragraph(f"Nickname: {c['nickname']}")
        if c.get('anniversary'):
            document.add_paragraph(f"Anniversary: {c['anniversary']}")
        if c.get('gender'):
            document.add_paragraph(f"Gender: {c['gender']}")
        if c.get('categories'):
            document.add_paragraph(f"Categories: {c['categories']}")
        for field_name, field_value in c.get('_custom', {}).items():
            document.add_paragraph(f"{field_name}: {field_value}")

        if c.get('_skipped'):
            for msg in c['_skipped']:
                document.add_paragraph(f"⚠ {msg}")

    document.save(str(out_path))


# ---------------------------------------------------------------------------
# Duplicate detection
# ---------------------------------------------------------------------------

def identity_key(contact):
    """
    Build a hashable identity key from fields that define who a person is.
    Metadata fields (revised, created, notes, categories, custom) are excluded.

    Tuple positions:
      0 — name
      1 — phones (bare values, no type labels)
      2 — emails (bare values, no type labels)
      3 — addresses
      4 — organizations
      5 — birthday

    If new identity fields are added to the script (e.g. IMPP),
    add them here too or duplicate detection will miss differences in that field.
    """
    def bare(entries):
        result = set()
        for e in entries:
            val = e.split(':', 1)[1].split('(')[0].strip() if ':' in e else e
            result.add(val.lower())
        return frozenset(result)

    return (
        (contact.get('name') or '').strip().lower(),
        bare(contact.get('phones', [])),
        bare(contact.get('emails', [])),
        bare(contact.get('addresses', [])),
        frozenset(o.lower() for o in contact.get('organizations', [])),
        (contact.get('birthday') or '').lower(),
    )


def merge_contacts(base, duplicate):
    """
    Merge a duplicate contact into the base contact.
    Identity fields (phones, emails, addresses, orgs) are already the same
    by definition — no need to merge those.

    Meaningful non-identity fields are merged if missing from base:
      title, nickname, anniversary, gender, categories, urls, label,
      role, timezone, geo (v4)

    IMPP is merged like urls — unique values from duplicate are added
    to base rather than all-or-nothing (a person can have multiple IM
    handles; this matches how multiple phones/emails/urls are handled).

    Notes and custom fields from duplicate are appended (not overwritten)
    with a warning logged so the user knows data came from a removed duplicate.

    Revised:  keep the newer timestamp  (most recent edit wins)
    Created:  keep the earlier timestamp (oldest origin wins)
    """
    # Fields: take from duplicate only if base is missing
    for field in ('title', 'nickname', 'anniversary', 'gender', 'categories', 'label',
                  'role', 'timezone', 'geo'):
        if not base.get(field) and duplicate.get(field):
            base[field] = duplicate[field]

    # IMPP: merge unique values, same pattern as urls below
    if duplicate.get('impp'):
        existing = set(base.get('impp', []))
        for handle in duplicate['impp']:
            if handle not in existing:
                base.setdefault('impp', []).append(handle)
                existing.add(handle)

    # URLs: merge unique values
    if duplicate.get('urls'):
        existing = set(base.get('urls', []))
        for url in duplicate['urls']:
            if url not in existing:
                base.setdefault('urls', []).append(url)
                existing.add(url)

    # Notes from duplicate appended with marker, skip if same content already exists
    if duplicate.get('notes'):
        base.setdefault('notes', [])
        existing_bare = {n.replace('[merged] ', '') for n in base['notes']}
        for note in duplicate['notes']:
            if note not in existing_bare:
                marked = f"[merged] {note}"
                if marked not in base['notes']:
                    base['notes'].append(marked)

    # Custom fields: merge unique keys, don't overwrite existing
    if duplicate.get('_custom'):
        base.setdefault('_custom', {})
        for key, val in duplicate['_custom'].items():
            if key not in base['_custom']:
                base['_custom'][key] = val

    # Revised: keep the newer timestamp (compare raw ISO strings, not formatted display strings)
    if duplicate.get('_revised_raw') and base.get('_revised_raw'):
        if duplicate['_revised_raw'] > base['_revised_raw']:
            base['revised'] = duplicate['revised']
            base['_revised_raw'] = duplicate['_revised_raw']
    elif duplicate.get('_revised_raw') and not base.get('_revised_raw'):
        base['revised'] = duplicate['revised']
        base['_revised_raw'] = duplicate['_revised_raw']

    # Created: keep the earlier timestamp — oldest origin date is the true one
    if duplicate.get('_created_raw') and base.get('_created_raw'):
        if duplicate['_created_raw'] < base['_created_raw']:
            base['created'] = duplicate['created']
            base['_created_raw'] = duplicate['_created_raw']
    elif duplicate.get('_created_raw') and not base.get('_created_raw'):
        base['created'] = duplicate['created']
        base['_created_raw'] = duplicate['_created_raw']

    return base


def is_newer(a, b):
    """
    Return True if contact a has a newer revised timestamp than b.
    Uses the raw vCard datetime string (_revised_raw) for comparison,
    not the formatted display string — ISO 8601 compares correctly as strings.
    If either is missing, returns False (keep existing base).
    """
    rev_a = a.get('_revised_raw') or ''
    rev_b = b.get('_revised_raw') or ''
    if not rev_a or not rev_b:
        return False
    return rev_a > rev_b


def classify_duplicates(contacts):
    """
    Two-tier duplicate classification:

    Exact  — all identity fields match → merge meaningful extra fields,
             newer REV contact becomes base, duplicate removed and reported.
    Fuzzy  — same name + at least one shared phone or email → kept, warned only.

    Returns (kept, exact_removed, fuzzy_warnings):
      kept          : list of contacts after exact merge/removal
      exact_removed : list of (retained_contact_number, name)
      fuzzy_warnings: list of (i, j, name), 1-based indices into kept
    """
    # key → index in kept list (for merging)
    seen = {}
    kept = []
    exact_removed = []

    for contact in contacts:
        key = identity_key(contact)
        if key in seen:
            kept_idx = seen[key]
            existing = kept[kept_idx]
            # REV-aware: newer contact becomes base, older is merged into it
            if is_newer(contact, existing):
                kept[kept_idx] = merge_contacts(contact, existing)
            else:
                kept[kept_idx] = merge_contacts(existing, contact)
            # Store (retained_contact_number, name) for warning output.
            # Fuzzy pass only warns, never removes, so this index stays accurate.
            exact_removed.append((kept_idx + 1, contact.get('name') or '(Unknown)'))
        else:
            seen[key] = len(kept)
            kept.append(contact)

    def phone_vals(c):
        return {
            p.split(':', 1)[1].split('(')[0].strip().lower()
            for p in c.get('phones', [])
        }

    def email_vals(c):
        return {
            e.split(':', 1)[1].split('(')[0].strip().lower()
            for e in c.get('emails', [])
        }

    fuzzy_warnings = []
    for i in range(len(kept)):
        for j in range(i + 1, len(kept)):
            a, b = kept[i], kept[j]
            name_a = (a.get('name') or '').strip().lower()
            name_b = (b.get('name') or '').strip().lower()
            if not name_a or not name_b or name_a != name_b:
                continue
            if phone_vals(a) & phone_vals(b) or email_vals(a) & email_vals(b):
                fuzzy_warnings.append((i + 1, j + 1, a.get('name', '(Unknown)')))

    return kept, exact_removed, fuzzy_warnings


# ---------------------------------------------------------------------------
# Preview
# ---------------------------------------------------------------------------

def run_paged_preview(contacts, fuzzy_pairs=None):
    """
    Paged terminal preview — zero dependencies, works everywhere.

    Displays contacts page by page with a prompt at the bottom.
    Each page shows PAGE_SIZE contacts with their index number.

    Controls:
      Enter / n   next page
      b           previous page
      /<term>     search by name, org, or phone (press Enter to clear)
      q           quit

    fuzzy_pairs is a set of 0-based contact indices flagged as fuzzy
    duplicates — shown with a '?' marker so the user can spot them.
    """
    PAGE_SIZE  = 8
    fuzzy_idx  = fuzzy_pairs or set()

    def _render_page(page_contacts, page_start, total_shown, query):
        """Print one page of contacts to the terminal."""
        # ANSI clear screen + move to top (works on Windows Terminal, VS Code, etc.)
        print('\033[2J\033[H', end='')
        print(f"vCardTra — Preview  ({total_shown} contacts"
              + (f", filtered: '{query}'" if query else '') + ")\n")
        print(f"  {'#':<5}  {'Name':<30}  {'Org':<25}  {'Phone'}")
        print(f"  {'-'*5}  {'-'*30}  {'-'*25}  {'-'*20}")
        for offset, (orig_idx, c) in enumerate(page_contacts):
            num     = page_start + offset + 1
            name    = (c.get('name') or '(Unknown)')[:30]
            org     = (c.get('organizations') or [''])[0][:25]
            phone   = ''
            if c.get('phones'):
                # Strip "Phone: " prefix and type label for compact display
                raw = c['phones'][0]
                phone = raw[len('Phone: '):].split(' (')[0][:20]
            marker = '?' if orig_idx in fuzzy_idx else ' '
            print(f"  {num:<5}{marker} {name:<30}  {org:<25}  {phone}")
        print()

    def _render_detail(contacts_window, page_start):
        """Print full detail of contacts on the current page."""
        print('\033[2J\033[H', end='')
        for offset, (_, c) in enumerate(contacts_window):
            num = page_start + offset + 1
            print(f"⭐ Contact {num}:")
            print('-' * 60)
            print(format_contact(c))
            print()
        input("  Press Enter to return to list...")

    # Build working list as (original_index, contact) pairs so fuzzy markers
    # stay correct even after filtering narrows the visible set
    active   = list(enumerate(contacts))
    query    = ''
    page     = 0

    while True:
        # Apply search filter
        if query:
            q = query.lower()
            visible = [
                (i, c) for i, c in active
                if q in (c.get('name') or '').lower()
                or q in ' '.join(c.get('organizations') or []).lower()
                or any(q in p.lower() for p in c.get('phones', []))
            ]
        else:
            visible = active

        total_pages = max(1, (len(visible) + PAGE_SIZE - 1) // PAGE_SIZE)
        page = max(0, min(page, total_pages - 1))

        page_slice = visible[page * PAGE_SIZE : (page + 1) * PAGE_SIZE]
        _render_page(page_slice, page * PAGE_SIZE, len(visible), query)

        # Prompt
        page_info = f"Page {page + 1}/{total_pages}"
        print(f"  {page_info}  |  Enter/n=next  b=back  d=detail  /search  q=quit")
        try:
            raw = input("  > ").strip()
        except (KeyboardInterrupt, EOFError):
            print()
            break

        if raw.lower() == 'q':
            break
        elif raw.lower() in ('', 'n'):
            page = (page + 1) % total_pages
        elif raw.lower() == 'b':
            page = (page - 1) % total_pages
        elif raw.lower() == 'd':
            _render_detail(page_slice, page * PAGE_SIZE)
        elif raw.startswith('/'):
            query = raw[1:].strip()
            page  = 0
        else:
            # Any other input — treat as implicit search if it looks like text
            pass

    print('\033[2J\033[H', end='')  # clear on exit


def run_textual_preview(contacts, fuzzy_pairs=None):
    """
    Full interactive TUI preview using the Textual library.

    Layout: two-panel — contact list (left) + contact detail (right).
    Keyboard shortcuts:
      /           open search bar; list filters live as you type
      Escape      close search bar / clear search
      space       select / deselect highlighted contact
      a           select all visible contacts
      n           clear all selections
      s           cycle sort: original → name A-Z → name Z-A → org A-Z → original
      e           export: selected contacts (or all if none selected) to a .txt file
      q / ctrl+c  quit

    Selected contacts shown with ● marker.
    Fuzzy duplicate pairs highlighted in yellow.

    Raises ImportError if textual is not installed — caller catches this
    and falls back to run_paged_preview().
    """
    from textual.app import App, ComposeResult
    from textual.widgets import Header, Footer, ListView, ListItem, Label, Static
    from textual.containers import Horizontal, Vertical
    from textual.binding import Binding
    from textual import events

    class ContactApp(App):
        """Two-panel vCard preview application."""

        CSS = """
        Screen { layout: horizontal; }
        #list-panel {
            width: 40%;
            border-right: solid $primary;
            overflow-y: auto;
        }
        #detail-panel {
            width: 60%;
            padding: 1 2;
            overflow-y: auto;
        }
        #search-bar {
            height: 1;
            background: $accent;
            color: $text;
            padding: 0 1;
            display: none;
        }
        #search-bar.visible { display: block; }
        #stats-bar {
            height: 1;
            background: $primary-darken-2;
            color: $text-muted;
            padding: 0 1;
            dock: bottom;
        }
        ListItem { padding: 0 1; }
        ListItem.selected { color: $success; }
        ListItem.fuzzy { color: $warning; }
        ListItem.selected.fuzzy { color: $warning; text-style: bold; }
        """

        BINDINGS = [
            Binding('q',      'quit',         'Quit'),
            Binding('/',      'search',        'Search'),
            Binding('escape', 'clear_search',  'Clear'),
            Binding('space',  'toggle_select', 'Select',  show=False),
            Binding('a',      'select_all',    'All'),
            Binding('n',      'clear_select',  'None'),
            Binding('s',      'cycle_sort',    'Sort'),
            Binding('e',      'export',        'Export'),
        ]

        def __init__(self, contacts, fuzzy_idx):
            super().__init__()
            self._all_contacts  = list(enumerate(contacts))  # (orig_idx, contact)
            self._fuzzy_idx     = fuzzy_idx
            self._selected      = set()   # orig_idx values
            self._query         = ''
            self._sort_mode     = 0       # 0=original 1=name↑ 2=name↓ 3=org↑
            self._sort_labels   = ['original', 'name A→Z', 'name Z→A', 'org A→Z']
            self._searching     = False
            self._visible       = list(self._all_contacts)

        def compose(self) -> ComposeResult:
            yield Header(show_clock=False)
            with Horizontal():
                with Vertical(id='list-panel'):
                    yield Static('', id='search-bar')
                    yield ListView(id='contact-list')
            yield Static('', id='detail-panel')
            yield Static('', id='stats-bar')
            yield Footer()

        def on_mount(self):
            self._refresh_list()
            self._update_stats()

        # ------------------------------------------------------------------ #
        # List management                                                      #
        # ------------------------------------------------------------------ #

        def _apply_query_and_sort(self):
            """Filter and sort _all_contacts into _visible."""
            result = self._all_contacts
            if self._query:
                q = self._query.lower()
                result = [
                    (i, c) for i, c in result
                    if q in (c.get('name') or '').lower()
                    or q in ' '.join(c.get('organizations') or []).lower()
                    or any(q in p.lower() for p in c.get('phones', []))
                ]
            if self._sort_mode == 1:
                result = sorted(result, key=lambda x: (not bool(x[1].get('name')), (x[1].get('name') or '').lower()))
            elif self._sort_mode == 2:
                known   = [x for x in result if x[1].get('name')]
                unknown = [x for x in result if not x[1].get('name')]
                result  = sorted(known, key=lambda x: x[1].get('name', '').lower(), reverse=True) + unknown
            elif self._sort_mode == 3:
                result = sorted(result, key=lambda x: (
                    not bool((x[1].get('organizations') or [''])[0]),
                    (x[1].get('organizations') or [''])[0].lower()
                ))
            self._visible = result

        def _refresh_list(self):
            """Rebuild the ListView from _visible."""
            self._apply_query_and_sort()
            lv = self.query_one('#contact-list', ListView)
            lv.clear()
            for pos, (orig_idx, c) in enumerate(self._visible):
                name    = c.get('name') or '(Unknown)'
                marker  = '●' if orig_idx in self._selected else ' '
                label   = f" {marker} {pos + 1:>4}.  {name}"
                item    = ListItem(Label(label))
                if orig_idx in self._selected:
                    item.add_class('selected')
                if orig_idx in self._fuzzy_idx:
                    item.add_class('fuzzy')
                lv.append(item)

        def _update_stats(self):
            bar = self.query_one('#stats-bar', Static)
            total   = len(self._all_contacts)
            visible = len(self._visible)
            sel     = len(self._selected)
            sort_l  = self._sort_labels[self._sort_mode]
            parts   = [f"{total} contacts"]
            if visible != total:
                parts.append(f"{visible} visible")
            if sel:
                parts.append(f"{sel} selected")
            parts.append(f"sort: {sort_l}")
            bar.update('  ' + '  ·  '.join(parts))

        def _show_detail(self, pos):
            """Render contact detail in the right panel."""
            panel = self.query_one('#detail-panel', Static)
            if pos < 0 or pos >= len(self._visible):
                panel.update('')
                return
            _, c = self._visible[pos]
            panel.update(format_contact(c))

        # ------------------------------------------------------------------ #
        # Event handlers                                                       #
        # ------------------------------------------------------------------ #

        def on_list_view_highlighted(self, event: ListView.Highlighted):
            if event.item is not None:
                idx = self.query_one('#contact-list', ListView).index
                self._show_detail(idx)

        def on_key(self, event: events.Key):
            if self._searching:
                if event.key == 'enter':
                    self._searching = False
                    sb = self.query_one('#search-bar', Static)
                    sb.remove_class('visible')
                    self._refresh_list()
                    self._update_stats()
                elif event.key == 'backspace':
                    self._query = self._query[:-1]
                    self.query_one('#search-bar', Static).update(f'/ {self._query}')
                    self._refresh_list()
                    self._update_stats()
                elif event.character and event.character.isprintable():
                    self._query += event.character
                    self.query_one('#search-bar', Static).update(f'/ {self._query}')
                    self._refresh_list()
                    self._update_stats()
                event.stop()

        # ------------------------------------------------------------------ #
        # Actions                                                              #
        # ------------------------------------------------------------------ #

        def action_search(self):
            self._searching = True
            sb = self.query_one('#search-bar', Static)
            sb.add_class('visible')
            sb.update(f'/ {self._query}')

        def action_clear_search(self):
            self._searching = False
            self._query     = ''
            sb = self.query_one('#search-bar', Static)
            sb.remove_class('visible')
            sb.update('')
            self._refresh_list()
            self._update_stats()

        def action_toggle_select(self):
            lv  = self.query_one('#contact-list', ListView)
            idx = lv.index
            if idx is None or idx >= len(self._visible):
                return
            orig_idx = self._visible[idx][0]
            if orig_idx in self._selected:
                self._selected.discard(orig_idx)
            else:
                self._selected.add(orig_idx)
            self._refresh_list()
            self._update_stats()
            # Restore highlight position after list rebuild
            lv.index = min(idx, len(self._visible) - 1)

        def action_select_all(self):
            for orig_idx, _ in self._visible:
                self._selected.add(orig_idx)
            self._refresh_list()
            self._update_stats()

        def action_clear_select(self):
            self._selected.clear()
            self._refresh_list()
            self._update_stats()

        def action_cycle_sort(self):
            self._sort_mode = (self._sort_mode + 1) % len(self._sort_labels)
            self._refresh_list()
            self._update_stats()

        def action_export(self):
            """
            Export contacts to a .txt file in the current directory.
            Exports selected contacts if any are selected, otherwise
            exports all currently visible (post-search/sort) contacts.
            """
            if self._selected:
                to_export = [c for orig_idx, c in self._all_contacts if orig_idx in self._selected]
                label = f"{len(to_export)} selected"
            else:
                to_export = [c for _, c in self._visible]
                label = f"{len(to_export)} visible"

            if not to_export:
                self.notify("Nothing to export.", severity="warning")
                return

            out_path = safe_output_path('preview_export.txt')
            lines = [
                "Generated by vCardTra — preview export",
                DIVIDER,
                f"  Exported  :  {datetime.now().strftime('%B %d, %Y, %H:%M')}",
                f"  Contacts  :  {label}",
                DIVIDER,
                "",
            ]
            for i, c in enumerate(to_export, 1):
                lines.append(f"⭐ Contact {i}:")
                lines.append(DIVIDER_CTX)
                lines.append(format_contact(c))
                lines.append("")

            try:
                Path(out_path).write_text('\n'.join(lines), encoding='utf-8')
                self.notify(f"Exported {label} → {out_path}", severity="information")
            except OSError as e:
                self.notify(f"Export failed: {e}", severity="error")

    # Build fuzzy pair index — collect all orig_idx values that appear in any pair
    fuzzy_set = set()
    if fuzzy_pairs:
        for i, j in fuzzy_pairs:
            fuzzy_set.add(i - 1)   # fuzzy_pairs uses 1-based indices
            fuzzy_set.add(j - 1)

    ContactApp(contacts, fuzzy_set).run()


def run_browser_preview(contacts, header_rows, single_mode, include_photos=True):
    """
    --preview --browser mode. Reuses format_html() (the same renderer
    --format html writes to disk) to build a self-contained HTML document,
    writes it to a temp file, and opens it in the system default browser
    via the stdlib webbrowser module — zero extra dependency, per the v4 plan.

    Unlike the TUI/paged preview modes, this is one-shot: no live
    search/select/export inside the browser. It's a quick "what would this
    look like" view, not an interactive session. Terminal preview remains
    the interactive option.
    """
    import webbrowser
    import tempfile

    html_content = format_html(contacts, header_rows, single_mode, include_photos=include_photos)
    with tempfile.NamedTemporaryFile('w', suffix='.html', delete=False, encoding='utf-8') as f:
        f.write(html_content)
        temp_path = f.name

    print(f"\n  Opening preview in your browser ({len(contacts)} contacts)...")
    print(f"  ({temp_path})")
    webbrowser.open(f"file://{temp_path}")


def run_preview(contacts, fuzzy_warnings=None):
    """
    Entry point for --preview mode.

    Tries to launch the Textual TUI first. If textual is not installed,
    falls back to the paged terminal preview silently — no error, no crash.

    fuzzy_warnings is the list of (i, j, name) tuples from classify_duplicates.
    We extract just the index pairs for highlighting.
    """
    # Build set of (i, j) 1-based pairs for highlighting
    fuzzy_pairs = {(i, j) for i, j, _ in (fuzzy_warnings or [])}

    try:
        run_textual_preview(contacts, fuzzy_pairs=fuzzy_pairs)
    except ImportError:
        # textual not installed — paged fallback, zero dependencies
        # Convert pairs to flat set of 0-based indices for the paged renderer
        fuzzy_idx = set()
        for i, j in fuzzy_pairs:
            fuzzy_idx.add(i - 1)
            fuzzy_idx.add(j - 1)
        run_paged_preview(contacts, fuzzy_pairs=fuzzy_idx)


# ---------------------------------------------------------------------------
# Converter
# ---------------------------------------------------------------------------

def convert(input_files, output_file, single_mode=False, sort_contacts=False, stats_only=False, encoding=None, sort_by=None, reverse=False, limit=None, filters=None, select=None, exclude=None, output_format='text', no_photos=False):
    """
    Parse one or more .vcf files and write all contacts to an output file.
    Auto-renames the output file if it already exists.

    encoding      : if set, passed to read_file() to bypass auto-detection.
    sort_by       : one of 'name', 'org', 'birthday', 'created' — implies sorting.
    reverse       : reverse the sort order.
    limit         : keep only the first N contacts after all filtering/selection.
    filters       : list of 'field=value' strings — all must match (AND logic).
    select        : selection expression string — keep only matched contacts.
    exclude       : selection expression string — remove matched contacts.
    output_format : one of FORMAT_EXTENSIONS keys ('text', 'json', ...). Only
                    the final write step branches on this — parsing, sort,
                    dedup, filter, select, and limit are format-agnostic.
    no_photos     : if True, PHOTO/LOGO fields are skipped during parsing
                    (faster; also means include_photos=False for JSON).

    Order of operations: parse → sort → dedup → filter → select/exclude → limit.
    """
    all_contacts = []
    malformed = 0
    total_found = 0

    for path_str in input_files:
        path = Path(path_str)
        if not path.exists():
            print(f"  ! File not found: {path_str}")
            continue

        print(f"  Reading {path.name} ...", end=' ', flush=True)

        try:
            content = read_file(path, encoding=encoding)
        except ValueError as e:
            print(f"\n  ! {e}")
            return False

        blocks = re.split(r'END:VCARD\s*', content, flags=re.IGNORECASE)
        blocks = [b + 'END:VCARD' for b in blocks if re.search(r'BEGIN:VCARD', b, re.IGNORECASE)]

        file_contacts = []
        file_malformed = 0
        for block in blocks:
            contact = parse_vcard(block, no_photos=no_photos)
            if contact and has_displayable_content(contact):
                contact['_source'] = path.name
                file_contacts.append(contact)
            else:
                file_malformed += 1

        total_found += len(file_contacts) + file_malformed
        malformed += file_malformed

        n_contacts = len(file_contacts)
        n_malformed = file_malformed
        status = f"{n_contacts} {'contact' if n_contacts == 1 else 'contacts'}"
        if n_malformed:
            status += f", {n_malformed} malformed"
        print(status)

        all_contacts.extend(file_contacts)

    if not all_contacts:
        print("\n  No contacts found.")
        if malformed:
            m = malformed
            print(f"  {m} malformed {'block' if m == 1 else 'blocks'} skipped.")
        return False

    # Sort — skip entirely in stats-only mode (order is never shown)
    do_sort = (sort_contacts or sort_by) and not stats_only
    if do_sort:
        # Sort key functions per field
        def _sort_key(c):
            if sort_by == 'org':
                primary = (c.get('organizations') or [''])[0].lower()
            elif sort_by == 'birthday':
                primary = (c.get('birthday') or '').lower()
            elif sort_by == 'created':
                primary = (c.get('_created_raw') or '').lower()
            else:  # 'name' or default --sort
                primary = (c.get('name') or '').lower()
            # unknowns (empty string) always sort last regardless of reverse
            return (not bool(primary), primary)

        all_contacts.sort(key=_sort_key, reverse=False)
        if reverse:
            # Reverse only the non-empty entries; keep unknowns at the end
            unknowns = [c for c in all_contacts if not _sort_key(c)[1]]
            known    = [c for c in all_contacts if _sort_key(c)[1]]
            all_contacts = known[::-1] + unknowns

    all_contacts, exact_removed, fuzzy_warnings = classify_duplicates(all_contacts)

    # Filter — applied after dedup+sort so contacts are in final display order.
    # Stats-only skips writing but still applies filters so the count is accurate.
    filtered_count = None  # None means no filter was active
    if filters:
        before = len(all_contacts)
        try:
            all_contacts = apply_filters(all_contacts, filters)
        except ValueError as e:
            print(f"\n  ! {e}")
            return False
        filtered_count = before - len(all_contacts)

    # Select / Exclude — operates on 1-based indices of the post-filter list.
    # Selection expression is validated inside apply_selection(); index errors
    # are caught here and reported cleanly.
    selected_expr = None
    excluded_expr = None
    if select:
        selected_expr = select
        try:
            all_contacts = apply_selection(all_contacts, select, exclude=False)
        except ValueError as e:
            print(f"\n  ! {e}")
            return False
    elif exclude:
        excluded_expr = exclude
        try:
            all_contacts = apply_selection(all_contacts, exclude, exclude=True)
        except ValueError as e:
            print(f"\n  ! {e}")
            return False

    # Limit — applied last, after dedup, filter, and select
    limited = False
    if limit and len(all_contacts) > limit:
        all_contacts = all_contacts[:limit]
        limited = True

    # Collect stats (used by both --stats mode and full output)
    total_skipped_fields = sum(len(c.get('_skipped', [])) for c in all_contacts)
    contacts_with_skips  = sum(1 for c in all_contacts if c.get('_skipped'))
    versions_found       = sorted({c.get('_version') for c in all_contacts if c.get('_version')})

    contacts_parts = [f"{total_found} found"]
    if malformed:
        contacts_parts.append(f"{malformed} malformed")
    if exact_removed:
        n_removed = len(exact_removed)
        contacts_parts.append(f"{n_removed} {'duplicate' if n_removed == 1 else 'duplicates'} removed")
    if filtered_count:
        contacts_parts.append(f"{filtered_count} filtered out")
    if selected_expr:
        contacts_parts.append("selection applied")
    if excluded_expr:
        contacts_parts.append("exclusion applied")
    if limited:
        contacts_parts.append(f"limited to {limit}")
    contacts_parts.append(f"{len(all_contacts)} exported")
    contacts_line = ', '.join(contacts_parts)

    # Stats-only mode: print summary and exit without writing a file
    if stats_only:
        print(f"\n  {DIVIDER}")
        print(f"  Contacts    :  {contacts_line}")
        if versions_found:
            print(f"  vCard       :  {', '.join(versions_found)}")
        if encoding:
            print(f"  Encoding    :  {encoding} (manual override)")
        if total_skipped_fields:
            fields_str   = f"{total_skipped_fields} {'field' if total_skipped_fields == 1 else 'fields'}"
            contacts_str = f"{contacts_with_skips} {'contact' if contacts_with_skips == 1 else 'contacts'}"
            print(f"  Skipped     :  {fields_str} across {contacts_str}")
        if fuzzy_warnings:
            w = len(fuzzy_warnings)
            print(f"  Warnings    :  {w} duplicate {'warning' if w == 1 else 'warnings'}")
        if do_sort:
            sort_label = sort_by or 'name'
            direction  = 'Z → A' if reverse else 'A → Z'
            print(f"  Sorted      :  {sort_label} {direction}")
        if filters:
            print(f"  Filter      :  {', '.join(filters)}")
        if selected_expr:
            print(f"  Select      :  {selected_expr}")
        if excluded_expr:
            print(f"  Exclude     :  {excluded_expr}")
        if limited:
            print(f"  Limit       :  {limit} contacts")
        print(f"  {DIVIDER}")
        return True

    out_path = safe_output_path(output_file)
    if out_path != Path(output_file):
        print(f"\n  '{Path(output_file).name}' already exists — saving as '{out_path.name}'")

    # Build source line
    if single_mode:
        source_label = 'Source'
        source_value = Path(input_files[0]).name
    else:
        source_label = 'Sources'
        names = [Path(f).name for f in input_files]
        if len(names) <= 3:
            source_value = ', '.join(names)
        else:
            source_value = f"{', '.join(names[:2])}, and {len(names) - 2} more"

    # Assemble header
    now = datetime.now().strftime('%B %d, %Y, %H:%M')
    header_rows = [
        ('Exported',  now),
        (source_label, source_value),
    ]
    if not single_mode:
        n = len(input_files)
        header_rows.append(('Files', f"{n} {'file' if n == 1 else 'files'}"))
    header_rows.append(('Contacts', contacts_line))
    if versions_found:
        header_rows.append(('vCard versions', ', '.join(versions_found)))
    if encoding:
        header_rows.append(('Encoding', f"{encoding} (manual override)"))
    if total_skipped_fields:
        sf = total_skipped_fields
        cs = contacts_with_skips
        skip_str = f"{sf} {'field' if sf == 1 else 'fields'} across {cs} {'contact' if cs == 1 else 'contacts'}"
        header_rows.append(('Skipped fields', skip_str))
    if fuzzy_warnings:
        header_rows.append(('Duplicate warnings', f"{len(fuzzy_warnings)} (review in warnings section)"))
    if do_sort:
        sort_label = sort_by or 'name'
        direction  = 'Z → A' if reverse else 'A → Z'
        header_rows.append(('Sorted', f"{sort_label} {direction}"))
    if filters:
        header_rows.append(('Filter', ', '.join(filters)))
    if selected_expr:
        header_rows.append(('Select', selected_expr))
    if excluded_expr:
        header_rows.append(('Exclude', excluded_expr))
    if limited:
        header_rows.append(('Limit', f"{limit} contacts"))

    # meta mirrors header_rows as native types/snake_case keys — built
    # unconditionally (not just for --format json) so XML can reuse it too.
    meta = {
        'generated_by': 'vCardTra',
        'version': __version__,
        'exported': now,
    }
    if single_mode:
        meta['source'] = source_value
    else:
        meta['sources'] = [Path(f).name for f in input_files]
        meta['file_count'] = len(input_files)
    meta['contacts_found'] = total_found
    if malformed:
        meta['malformed'] = malformed
    if exact_removed:
        meta['duplicates_removed'] = len(exact_removed)
        meta['duplicate_merges'] = [
            {'name': name, 'merged_into_contact': idx} for idx, name in exact_removed
        ]
    if filtered_count:
        meta['filtered_out'] = filtered_count
    if selected_expr:
        meta['select'] = selected_expr
    if excluded_expr:
        meta['exclude'] = excluded_expr
    if limited:
        meta['limit'] = limit
    meta['contacts_exported'] = len(all_contacts)
    if versions_found:
        meta['vcard_versions'] = versions_found
    if encoding:
        meta['encoding_override'] = encoding
    if total_skipped_fields:
        meta['skipped_fields_total'] = total_skipped_fields
        meta['contacts_with_skipped_fields'] = contacts_with_skips
    if fuzzy_warnings:
        meta['duplicate_warnings'] = [
            {'contact_a': i, 'contact_b': j, 'name': name} for i, j, name in fuzzy_warnings
        ]
    if do_sort:
        meta['sorted_by'] = sort_by or 'name'
        meta['sort_direction'] = 'desc' if reverse else 'asc'

    # Format dispatch. Most formatters return a string written via write_text;
    # 'sqlite', 'pdf', and 'docx' write the file themselves (binary / library-
    # owned save), so those three are handled as a direct call instead.
    include_photos = not no_photos
    try:
        if output_format == 'text':
            out_path.write_text(
                _render_text(all_contacts, header_rows, single_mode, malformed, exact_removed, fuzzy_warnings),
                encoding='utf-8')
        elif output_format == 'json':
            out_path.write_text(format_json(all_contacts, meta, include_photos=include_photos), encoding='utf-8')
        elif output_format == 'csv':
            out_path.write_text(format_csv(all_contacts), encoding='utf-8')
        elif output_format in ('markdown', 'md'):
            out_path.write_text(format_markdown(all_contacts, header_rows, single_mode), encoding='utf-8')
        elif output_format == 'html':
            out_path.write_text(
                format_html(all_contacts, header_rows, single_mode, include_photos=include_photos),
                encoding='utf-8')
        elif output_format == 'ics':
            ics_text, ics_skipped = format_ics(all_contacts)
            out_path.write_text(ics_text, encoding='utf-8')
            if ics_skipped:
                print(f"  ! {ics_skipped} birthday/anniversary date(s) had no day-level precision — skipped in the calendar")
        elif output_format == 'sql':
            out_path.write_text(format_sql(all_contacts), encoding='utf-8')
        elif output_format == 'sqlite':
            write_sqlite(all_contacts, out_path)
        elif output_format == 'xml':
            out_path.write_text(format_xml(all_contacts, meta, include_photos=include_photos), encoding='utf-8')
        elif output_format == 'pdf':
            write_pdf(all_contacts, header_rows, single_mode, out_path, include_photos=include_photos)
        elif output_format == 'docx':
            write_docx(all_contacts, header_rows, single_mode, out_path, include_photos=include_photos)
    except OSError as e:
        print(f"\n  ! Could not write output file: {e}")
        return False
    except ImportError as e:
        print(f"\n  ! {e}")
        return False

    # Terminal summary
    print(f"\n  {DIVIDER}")
    print(f"  Contacts    :  {contacts_line}")
    if versions_found:
        print(f"  vCard       :  {', '.join(versions_found)}")
    if encoding:
        print(f"  Encoding    :  {encoding} (manual override)")
    if total_skipped_fields:
        fields_str   = f"{total_skipped_fields} {'field' if total_skipped_fields == 1 else 'fields'}"
        contacts_str = f"{contacts_with_skips} {'contact' if contacts_with_skips == 1 else 'contacts'}"
        print(f"  Skipped     :  {fields_str} across {contacts_str}")
    if fuzzy_warnings:
        w = len(fuzzy_warnings)
        print(f"  Warnings    :  {w} duplicate {'warning' if w == 1 else 'warnings'} — review output file")
    if do_sort:
        sort_label = sort_by or 'name'
        direction  = 'Z → A' if reverse else 'A → Z'
        print(f"  Sorted      :  {sort_label} {direction}")
    if filters:
        print(f"  Filter      :  {', '.join(filters)}")
    if selected_expr:
        print(f"  Select      :  {selected_expr}")
    if excluded_expr:
        print(f"  Exclude     :  {excluded_expr}")
    if limited:
        print(f"  Limit       :  {limit} contacts")
    print(f"  Output      :  {out_path}")
    print(f"  {DIVIDER}")

    return True


def _load_contacts(input_files, encoding=None):
    """
    Shared file-reading logic for convert_merge() and convert_split().
    Returns (all_contacts, total_found, malformed).
    """
    all_contacts = []
    malformed = 0
    total_found = 0

    for path_str in input_files:
        path = Path(path_str)
        if not path.exists():
            print(f"  ! File not found: {path_str}")
            continue

        print(f"  Reading {path.name} ...", end=' ', flush=True)

        try:
            content = read_file(path, encoding=encoding)
        except ValueError as e:
            print(f"\n  ! {e}")
            return None, 0, 0

        blocks = re.split(r'END:VCARD\s*', content, flags=re.IGNORECASE)
        blocks = [b + 'END:VCARD' for b in blocks if re.search(r'BEGIN:VCARD', b, re.IGNORECASE)]

        file_contacts = []
        file_malformed = 0
        for block in blocks:
            contact = parse_vcard(block)
            if contact and has_displayable_content(contact):
                contact['_source'] = path.name
                file_contacts.append(contact)
            else:
                file_malformed += 1

        total_found += len(file_contacts) + file_malformed
        malformed += file_malformed

        n_c = len(file_contacts)
        status = f"{n_c} {'contact' if n_c == 1 else 'contacts'}"
        if file_malformed:
            status += f", {file_malformed} malformed"
        print(status)

        all_contacts.extend(file_contacts)

    return all_contacts, total_found, malformed


def convert_merge(input_files, output_file, encoding=None):
    """
    Merge one or more .vcf files into a single clean .vcf output.
    Runs full duplicate detection — exact duplicates merged and removed,
    fuzzy duplicates kept with a warning printed to terminal.
    Output is always vCard 3.0.
    """
    all_contacts, total_found, malformed = _load_contacts(input_files, encoding=encoding)

    if all_contacts is None:
        return False

    if not all_contacts:
        print("\n  No contacts found.")
        if malformed:
            m = malformed
            print(f"  {m} malformed {'block' if m == 1 else 'blocks'} skipped.")
        return False

    all_contacts, exact_removed, fuzzy_warnings = classify_duplicates(all_contacts)

    out_path = safe_output_path(output_file)
    if out_path != Path(output_file):
        print(f"\n  '{Path(output_file).name}' already exists — saving as '{out_path.name}'")

    blocks = [format_vcf(c) for c in all_contacts]
    try:
        out_path.write_text('\n\n'.join(blocks) + '\n', encoding='utf-8')
    except OSError as e:
        print(f"\n  ! Could not write output file: {e}")
        return False

    # Terminal summary
    contacts_parts = [f"{total_found} found"]
    if malformed:
        contacts_parts.append(f"{malformed} malformed")
    if exact_removed:
        n = len(exact_removed)
        contacts_parts.append(f"{n} {'duplicate' if n == 1 else 'duplicates'} removed")
    contacts_parts.append(f"{len(all_contacts)} exported")
    contacts_line = ', '.join(contacts_parts)

    print(f"\n  {DIVIDER}")
    print(f"  Contacts    :  {contacts_line}")
    if exact_removed:
        for retained_idx, name in exact_removed:
            print(f"                   - {name} → merged into contact {retained_idx}")
    if fuzzy_warnings:
        w = len(fuzzy_warnings)
        print(f"  Warnings    :  {w} possible {'duplicate' if w == 1 else 'duplicates'} — same name + shared phone/email")
        for i, j, name in fuzzy_warnings:
            print(f"                   - Contact {i} and Contact {j} ({name})")
    if encoding:
        print(f"  Encoding    :  {encoding} (manual override)")
    print(f"  Output      :  {out_path}")
    print(f"  {DIVIDER}")

    return True


def convert_split(input_files, output_dir, encoding=None):
    """
    Split one or more .vcf files into individual per-contact .vcf files.
    Runs full duplicate detection before splitting — exact duplicates are
    merged and removed so the output set is clean.
    Output files are named contacts_001.vcf, contacts_002.vcf, etc.,
    written into output_dir (default: current directory).
    """
    all_contacts, total_found, malformed = _load_contacts(input_files, encoding=encoding)

    if all_contacts is None:
        return False

    if not all_contacts:
        print("\n  No contacts found.")
        if malformed:
            m = malformed
            print(f"  {m} malformed {'block' if m == 1 else 'blocks'} skipped.")
        return False

    all_contacts, exact_removed, fuzzy_warnings = classify_duplicates(all_contacts)

    out_dir = Path(output_dir)
    try:
        out_dir.mkdir(parents=True, exist_ok=True)
    except OSError as e:
        print(f"\n  ! Could not create output directory: {e}")
        return False

    width = len(str(len(all_contacts)))  # zero-pad width based on total count
    written = 0
    for i, contact in enumerate(all_contacts, 1):
        filename = f"contact_{str(i).zfill(width)}.vcf"
        out_path = out_dir / filename
        try:
            out_path.write_text(format_vcf(contact) + '\n', encoding='utf-8')
            written += 1
        except OSError as e:
            print(f"  ! Could not write {filename}: {e}")

    # Terminal summary
    contacts_parts = [f"{total_found} found"]
    if malformed:
        contacts_parts.append(f"{malformed} malformed")
    if exact_removed:
        n = len(exact_removed)
        contacts_parts.append(f"{n} {'duplicate' if n == 1 else 'duplicates'} removed")
    contacts_parts.append(f"{written} exported")
    contacts_line = ', '.join(contacts_parts)

    print(f"\n  {DIVIDER}")
    print(f"  Contacts    :  {contacts_line}")
    if exact_removed:
        for retained_idx, name in exact_removed:
            print(f"                   - {name} → merged into contact {retained_idx}")
    if fuzzy_warnings:
        w = len(fuzzy_warnings)
        print(f"  Warnings    :  {w} possible {'duplicate' if w == 1 else 'duplicates'} — same name + shared phone/email")
        for i, j, name in fuzzy_warnings:
            print(f"                   - Contact {i} and Contact {j} ({name})")
    if encoding:
        print(f"  Encoding    :  {encoding} (manual override)")
    print(f"  Output      :  {written} files in {out_dir}/")
    print(f"  {DIVIDER}")

    return True


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    if '--version' in sys.argv:
        print(f"vCardTra {__version__}")
        sys.exit(0)

    if len(sys.argv) < 2 or '--help' in sys.argv or '-h' in sys.argv:
        print("vCardTra — Convert vCard files to text, JSON, and other formats")
        print("=" * 50)
        print("Usage:")
        print("  python vCardTra.py <file.vcf> [options]")
        print("  python vCardTra.py *.vcf -o contacts.txt")
        print()
        print("Options:")
        print("  -o <file>              Output file or directory")
        print("  --format <name>        Output format: text (default), json")
        print("  --no-photos            Skip decoding PHOTO/LOGO fields (faster; omits them from JSON)")
        print("  --sort                 Sort contacts A to Z by name")
        print("  --sort-by <field>      Sort by field: name, org, birthday, created")
        print("  --reverse              Reverse sort order (requires --sort or --sort-by)")
        print("  --limit <n>            Export only the first N contacts")
        print("  --filter <cond>        Filter contacts (repeatable, AND logic)")
        print("                           name=John   org=Acme   category=Work")
        print("                           has=phone   has=email  has=birthday")
        print("                           has=address has=note   has=url")
        print("  --select <expr>        Keep only matched contacts")
        print("                           1,5,10      1-10       1-10,15,20-25")
        print("                           last-10     John*      (name wildcard)")
        print("  --exclude <expr>       Remove matched contacts (same syntax as --select)")
        print("  --stats                Show summary statistics only, no output file written")
        print("  --preview              Interactive preview (TUI if textual installed, paged fallback)")
        print("  --browser              With --preview: open a one-shot HTML view in your browser instead")
        print("  --merge                Merge input VCFs into one clean .vcf file")
        print("  --split                Split contacts into individual .vcf files")
        print("  --encoding <codec>     Force a specific file encoding (e.g. shift-jis, gb2312)")
        print("  --version              Show program version")
        print("  -h, --help             Show this help")
        print()
        print("Examples:")
        print("  python vCardTra.py contacts.vcf")
        print("  python vCardTra.py contacts.vcf --format json")
        print("  python vCardTra.py contacts.vcf --format json --no-photos")
        print("  python vCardTra.py contacts.vcf --sort -o out.txt")
        print("  python vCardTra.py contacts.vcf --sort-by birthday")
        print("  python vCardTra.py contacts.vcf --sort --reverse --limit 10")
        print("  python vCardTra.py contacts.vcf --filter org=Acme --filter has=email")
        print("  python vCardTra.py contacts.vcf --sort --select 1-20")
        print("  python vCardTra.py contacts.vcf --exclude last-5")
        print("  python vCardTra.py contacts.vcf --preview")
        print("  python vCardTra.py contacts.vcf --sort --filter has=phone --preview")
        print("  python vCardTra.py contacts.vcf --stats")
        print("  python vCardTra.py contacts.vcf --encoding shift-jis")
        print("  python vCardTra.py file1.vcf file2.vcf --merge -o merged.vcf")
        print("  python vCardTra.py contacts.vcf --split -o split_output/")
        sys.exit(0 if '--help' in sys.argv or '-h' in sys.argv else 1)

    args = sys.argv[1:]
    output_file   = None
    sort_contacts = '--sort' in args
    stats_only    = '--stats' in args
    do_merge      = '--merge' in args
    do_split      = '--split' in args
    do_preview    = '--preview' in args
    do_browser    = '--browser' in args
    reverse       = '--reverse' in args
    no_photos     = '--no-photos' in args
    encoding      = None
    sort_by       = None
    limit         = None
    filters       = []
    select        = None
    exclude       = None
    output_format = 'text'

    if do_merge and do_split:
        print("Error: --merge and --split cannot be used together.")
        sys.exit(1)

    if do_preview and (do_merge or do_split or stats_only):
        print("Error: --preview cannot be combined with --merge, --split, or --stats.")
        sys.exit(1)

    if do_browser and not do_preview:
        print("Error: --browser requires --preview (it selects the preview's display mode).")
        sys.exit(1)

    args = [a for a in args if a not in ('--sort', '--stats', '--merge', '--split', '--reverse', '--preview', '--no-photos', '--browser')]

    if '--format' in args:
        idx = args.index('--format')
        if idx + 1 < len(args):
            output_format = args[idx + 1].lower()
            args = args[:idx] + args[idx + 2:]
        else:
            print("Error: --format requires a format name (e.g. --format json).")
            sys.exit(1)
        if output_format not in FORMAT_EXTENSIONS:
            if output_format in PLANNED_FORMATS:
                print(f"Error: --format {output_format} is planned but not yet implemented.")
                print(f"Available now: {', '.join(sorted(FORMAT_EXTENSIONS))}")
            else:
                print(f"Error: Unknown format '{output_format}'.")
                print(f"Available now: {', '.join(sorted(FORMAT_EXTENSIONS))}")
            sys.exit(1)
        if output_format in FORMAT_DEPENDENCIES:
            import importlib.util
            pkg_import_name = 'weasyprint' if output_format == 'pdf' else 'docx'
            if importlib.util.find_spec(pkg_import_name) is None:
                dep = FORMAT_DEPENDENCIES[output_format]
                print(f"Error: --format {output_format} requires the '{dep}' package, which isn't installed.")
                print(f"  Install with: pip install {dep}")
                sys.exit(1)
        if output_format != 'text' and (do_merge or do_split):
            print("Error: --format cannot be combined with --merge or --split — those always write valid .vcf files.")
            sys.exit(1)
        if output_format != 'text' and do_preview:
            print("Error: --format cannot be combined with --preview — preview always exports a text snapshot.")
            sys.exit(1)

    if '--sort-by' in args:
        idx = args.index('--sort-by')
        if idx + 1 < len(args):
            sort_by = args[idx + 1].lower()
            if sort_by not in ('name', 'org', 'birthday', 'created'):
                print("Error: --sort-by must be one of: name, org, birthday, created")
                sys.exit(1)
            args = args[:idx] + args[idx + 2:]
        else:
            print("Error: --sort-by requires a field name (name, org, birthday, created).")
            sys.exit(1)

    if '--limit' in args:
        idx = args.index('--limit')
        if idx + 1 < len(args):
            try:
                limit = int(args[idx + 1])
                if limit < 1:
                    raise ValueError
            except ValueError:
                print("Error: --limit requires a positive integer.")
                sys.exit(1)
            args = args[:idx] + args[idx + 2:]
        else:
            print("Error: --limit requires a number.")
            sys.exit(1)

    # --filter is repeatable; collect all occurrences
    while '--filter' in args:
        idx = args.index('--filter')
        if idx + 1 < len(args):
            filters.append(args[idx + 1])
            args = args[:idx] + args[idx + 2:]
        else:
            print("Error: --filter requires a condition (e.g. --filter org=Acme).")
            sys.exit(1)

    if '--select' in args:
        idx = args.index('--select')
        if idx + 1 < len(args):
            select = args[idx + 1]
            args = args[:idx] + args[idx + 2:]
        else:
            print("Error: --select requires an expression (e.g. --select 1-10).")
            sys.exit(1)

    if '--exclude' in args:
        idx = args.index('--exclude')
        if idx + 1 < len(args):
            exclude = args[idx + 1]
            args = args[:idx] + args[idx + 2:]
        else:
            print("Error: --exclude requires an expression (e.g. --exclude 1,5,9).")
            sys.exit(1)

    if select and exclude:
        print("Error: --select and --exclude cannot be used together.")
        sys.exit(1)

    if reverse and not sort_contacts and not sort_by:
        print("Error: --reverse requires --sort or --sort-by.")
        sys.exit(1)

    if '--encoding' in args:
        idx = args.index('--encoding')
        if idx + 1 < len(args):
            encoding = args[idx + 1]
            args = args[:idx] + args[idx + 2:]
        else:
            print("Error: --encoding flag requires a codec name (e.g. --encoding shift-jis).")
            sys.exit(1)

    if '-o' in args:
        idx = args.index('-o')
        if idx + 1 < len(args):
            output_file = args[idx + 1]
            args = args[:idx] + args[idx + 2:]
        else:
            print("Error: -o flag requires a filename.")
            sys.exit(1)

    if not args:
        print("Error: No input files specified.")
        sys.exit(1)

    input_files = []
    for pattern in args:
        matches = glob.glob(pattern)
        if matches:
            input_files.extend(matches)
        elif Path(pattern).exists():
            input_files.append(pattern)
        else:
            print(f"  ! No files matched: {pattern}")

    input_files = [f for f in input_files if f.lower().endswith('.vcf')]

    if not input_files:
        print("Error: No .vcf files found.")
        sys.exit(1)

    single_mode = len(input_files) == 1

    n = len(input_files)
    print(f"\nvCardTra — {n} {'file' if n == 1 else 'files'} to process\n")

    if do_merge:
        if output_file is None:
            names = [Path(f).stem for f in input_files]
            if len(names) == 1:
                output_file = f"{names[0]}_merged.vcf"
            elif len(names) == 2:
                output_file = f"{names[0]}_{names[1]}.vcf"
            else:
                output_file = f"{names[0]}_and_{len(names) - 1}_more.vcf"
        if not convert_merge(input_files, output_file, encoding=encoding):
            sys.exit(1)

    elif do_split:
        if output_file is None:
            output_file = f"split_{Path(input_files[0]).stem}"
        if not convert_split(input_files, output_file, encoding=encoding):
            sys.exit(1)

    elif do_preview:
        # Preview: run full parse → dedup → sort → filter → select/exclude pipeline,
        # then hand the resulting contact list to run_preview() instead of writing a file.
        # All sorting/filtering flags work exactly as in normal mode.
        contacts, total_found, malformed = _load_contacts(input_files, encoding=encoding)
        if not contacts:
            if malformed:
                m = malformed
                print(f"  {m} malformed {'block' if m == 1 else 'blocks'} skipped.")
            else:
                print("\n  No contacts found.")
            sys.exit(1)

        do_sort = sort_contacts or sort_by
        if do_sort:
            def _sort_key(c):
                if sort_by == 'org':
                    primary = (c.get('organizations') or [''])[0].lower()
                elif sort_by == 'birthday':
                    primary = (c.get('birthday') or '').lower()
                elif sort_by == 'created':
                    primary = (c.get('_created_raw') or '').lower()
                else:
                    primary = (c.get('name') or '').lower()
                return (not bool(primary), primary)
            contacts.sort(key=_sort_key)
            if reverse:
                unknowns = [c for c in contacts if not _sort_key(c)[1]]
                known    = [c for c in contacts if _sort_key(c)[1]]
                contacts = known[::-1] + unknowns

        contacts, exact_removed, fuzzy_warnings = classify_duplicates(contacts)

        if filters:
            try:
                contacts = apply_filters(contacts, filters)
            except ValueError as e:
                print(f"\n  ! {e}")
                sys.exit(1)

        if select:
            try:
                contacts = apply_selection(contacts, select, exclude=False)
            except ValueError as e:
                print(f"\n  ! {e}")
                sys.exit(1)
        elif exclude:
            try:
                contacts = apply_selection(contacts, exclude, exclude=True)
            except ValueError as e:
                print(f"\n  ! {e}")
                sys.exit(1)

        if limit and len(contacts) > limit:
            contacts = contacts[:limit]

        if not contacts:
            print("\n  No contacts to preview after filtering.")
            sys.exit(1)

        if do_browser:
            header_rows = [
                ('Exported', datetime.now().strftime('%B %d, %Y, %H:%M')),
                ('Source' if single_mode else 'Sources', ', '.join(Path(f).name for f in input_files)),
                ('Contacts', f"{len(contacts)} shown"),
            ]
            run_browser_preview(contacts, header_rows, single_mode, include_photos=not no_photos)
        else:
            run_preview(contacts, fuzzy_warnings=fuzzy_warnings)

    else:
        ext = FORMAT_EXTENSIONS[output_format]
        if output_file is None:
            if single_mode:
                output_file = str(Path(input_files[0]).with_suffix(f'.{ext}'))
            else:
                names = [Path(f).stem for f in input_files]
                output_file = (
                    f"{names[0]}_{names[1]}.{ext}" if len(names) == 2
                    else f"{names[0]}_and_{len(names) - 1}_more.{ext}"
                )
        if not convert(input_files, output_file, single_mode=single_mode, sort_contacts=sort_contacts, stats_only=stats_only, encoding=encoding, sort_by=sort_by, reverse=reverse, limit=limit, filters=filters, select=select, exclude=exclude, output_format=output_format, no_photos=no_photos):
            sys.exit(1)


if __name__ == "__main__":
    main()
